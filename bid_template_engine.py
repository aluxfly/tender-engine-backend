"""
标书 AI 生成系统 — Day 2: 动态模板引擎
========================================
功能:
  1. 根据招标项目类型自动匹配模板
  2. 从模板目录加载技术方案模板
  3. 生成三套模板文件（报价/商务/技术）
  4. 保存到 bid_templates 表
"""

import os
import re
import json
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 模板文件目录
TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

# 占位符正则：匹配 [xxx] 格式
PLACEHOLDER_RE = re.compile(r"\[([^\]]+)\]")

# ==================== 模板类型映射 ====================

# 根据招标关键词推断项目类型
PROJECT_TYPE_KEYWORDS = {
    "物联网卡": ["物联网卡", "SIM卡", "物联网", "M2M", "NB-IoT", "流量", "通信卡"],
    "布控球": ["布控球", "视频监控", "摄像头", "安防", "监控球", "球机"],
    "工程": ["工程", "施工", "建筑", "装修", "市政", "道路", "桥梁"],
    "软件": ["软件", "系统开发", "信息化", "数字化", "平台", "APP"],
    "设备": ["设备", "采购", "仪器", "硬件"],
    "服务": ["服务", "运维", "咨询", "培训", "监理"],
}

# 三套模板类型
TEMPLATE_TYPES = ["报价", "商务", "技术"]

# ==================== 占位符提取 ====================

def extract_placeholders_from_docx(doc: Document) -> List[Dict[str, Any]]:
    """从 docx Document 对象中提取所有占位符及其位置信息。"""
    placeholders = []
    seen = set()

    # 遍历段落
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        for match in PLACEHOLDER_RE.finditer(text):
            ph_name = match.group(1).strip()
            if ph_name and ph_name not in seen:
                seen.add(ph_name)
                placeholders.append({
                    "name": ph_name,
                    "location": f"paragraph_{para_idx}",
                    "context": text[max(0, match.start()-20):match.end()+20].strip(),
                })

    # 遍历表格
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text
                for match in PLACEHOLDER_RE.finditer(text):
                    ph_name = match.group(1).strip()
                    if ph_name and ph_name not in seen:
                        seen.add(ph_name)
                        placeholders.append({
                            "name": ph_name,
                            "location": f"table_{tbl_idx}_r{row_idx}_c{col_idx}",
                            "context": text[max(0, match.start()-20):match.end()+20].strip(),
                        })

    return placeholders


def classify_placeholder(name: str) -> str:
    """将占位符按类别分组。"""
    # 公司信息类
    company_keywords = ["公司", "企业", "名称", "地址", "注册", "信用代码", "统一社会", "经营范围",
                         "注册资本", "成立", "营业期限", "邮编", "传真"]
    # 财务类
    finance_keywords = ["财务", "金额", "报价", "单价", "合价", "大写", "小写", "预算", "保证金",
                         "税费", "审计", "资产", "利润", "纳税", "银行", "开户", "账号", "资信"]
    # 资质类
    qualification_keywords = ["资质", "证书", "执照", "许可证", "等级", "ISO", "认证",
                               "税务", "完税", "社保", "守法"]
    # 业绩类
    performance_keywords = ["业绩", "项目", "合同", "类似", "经验", "完成", "中标"]
    # 团队类
    team_keywords = ["人员", "团队", "经理", "负责人", "姓名", "职称", "学历", "专业",
                      "工作年限", "简历", "身份证", "授权", "代表", "法人", "签字",
                      "盖章", "被授权人", "职务"]
    # 技术类
    tech_keywords = ["技术", "方案", "偏差", "参数", "规格", "设备", "材料", "质量",
                      "安全", "风险", "进度", "验收", "培训", "售后", "服务",
                      "条款", "响应", "招标要求"]
    # 报价类
    price_keywords = ["报价", "价格", "单价", "合价", "总价", "费率", "折扣"]

    # 按优先级匹配
    if any(kw in name for kw in price_keywords):
        return "报价"
    if any(kw in name for kw in finance_keywords):
        return "财务"
    if any(kw in name for kw in qualification_keywords):
        return "资质"
    if any(kw in name for kw in performance_keywords):
        return "业绩"
    if any(kw in name for kw in team_keywords):
        return "团队"
    if any(kw in name for kw in tech_keywords):
        return "技术"
    if any(kw in name for kw in company_keywords):
        return "公司信息"

    return "其他"


# ==================== 模板加载 ====================

def load_default_template_docx(template_type: str) -> Optional[Document]:
    """从模板目录加载默认 docx 模板文件。"""
    type_file_map = {
        "报价": "bid-price-file.docx",
        "商务": "bid-business-file.docx",
        "技术": "bid-technical-file.docx",
    }
    filename = type_file_map.get(template_type)
    if not filename:
        return None

    filepath = TEMPLATES_DIR / filename
    if filepath.exists():
        return Document(str(filepath))

    # 尝试回退到已有的旧模板路径
    fallback = Path(__file__).parent / "templates" / filename
    if fallback.exists():
        return Document(str(fallback))

    return None


def generate_template_content(template_type: str, project_info: Dict[str, Any]) -> Dict[str, Any]:
    """
    生成单个模板的内容描述（JSON 格式存储到数据库）。

    返回:
    {
        "template_type": "报价/商务/技术",
        "template_name": "...",
        "file_path": "...",
        "placeholders": [...],  # 占位符列表
        "project_info": {...},  # 项目信息摘要
        "generated_at": "...",
    }
    """
    name_map = {
        "报价": f"报价文件_{project_info.get('title', '未知项目')}",
        "商务": f"商务文件_{project_info.get('title', '未知项目')}",
        "技术": f"技术文件_{project_info.get('title', '未知项目')}",
    }

    # 加载默认模板
    doc = load_default_template_docx(template_type)
    placeholders = []
    file_path = None

    if doc:
        placeholders = extract_placeholders_from_docx(doc)
        # 保存填充了项目信息的模板
        output_filename = f"template_{template_type}_{project_info.get('id', 'draft')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        file_path = str(TEMPLATES_DIR / output_filename)

        # 替换项目级占位符
        project_replacements = {
            "项目名称": project_info.get("project_name", "[项目名称]"),
            "招标编号": project_info.get("bid_number", "[招标编号]"),
        }
        for para in doc.paragraphs:
            for old, new in project_replacements.items():
                if old in para.text:
                    for run in para.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, str(new))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old, new in project_replacements.items():
                        if old in cell.text:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if old in run.text:
                                        run.text = run.text.replace(old, str(new))

        doc.save(file_path)
        logger.info(f"模板文件已保存: {file_path}")
    else:
        # 没有 docx 模板时，生成基础占位符列表
        file_path = None
        if template_type == "报价":
            placeholders = [
                {"name": "招标人名称", "location": "投标函", "context": "致：[招标人名称]"},
                {"name": "项目名称", "location": "投标函", "context": "[项目名称]项目"},
                {"name": "招标编号", "location": "投标函", "context": "招标编号：[招标编号]"},
                {"name": "授权代表姓名", "location": "投标函", "context": "[授权代表姓名]"},
                {"name": "人民币大写金额", "location": "价格表", "context": "[人民币大写金额]"},
                {"name": "小写金额", "location": "价格表", "context": "¥[小写金额]"},
                {"name": "投标人名称", "location": "签名", "context": "[投标人名称]"},
                {"name": "投标人地址", "location": "签名", "context": "[投标人地址]"},
                {"name": "联系电话", "location": "签名", "context": "[联系电话]"},
            ]
        elif template_type == "商务":
            placeholders = [
                {"name": "投标人名称", "location": "基本情况表", "context": "[投标人名称]"},
                {"name": "代码", "location": "基本情况表", "context": "[代码]"},
                {"name": "注册地址", "location": "基本情况表", "context": "[注册地址]"},
                {"name": "注册资本金额", "location": "基本情况表", "context": "[注册资本金额]"},
                {"name": "姓名", "location": "基本情况表/授权书", "context": "[姓名]"},
                {"name": "电话号码", "location": "基本情况表", "context": "[电话号码]"},
                {"name": "邮箱地址", "location": "基本情况表", "context": "[邮箱地址]"},
                {"name": "保证金大写金额", "location": "保证金", "context": "[保证金大写金额]"},
                {"name": "保证金小写金额", "location": "保证金", "context": "[保证金小写金额]"},
                {"name": "银行名称", "location": "保证金/财务", "context": "[银行名称]"},
                {"name": "银行账号", "location": "保证金", "context": "[银行账号]"},
            ]
        elif template_type == "技术":
            placeholders = [
                {"name": "条款号", "location": "技术偏差表", "context": "[条款号]"},
                {"name": "技术要求内容", "location": "技术偏差表", "context": "[技术要求内容]"},
                {"name": "响应内容", "location": "技术偏差表", "context": "[响应内容]"},
                {"name": "项目名称", "location": "类似项目表", "context": "[项目名称]"},
                {"name": "金额", "location": "类似项目表", "context": "[金额]"},
                {"name": "年", "location": "类似项目表", "context": "[年]"},
                {"name": "职称", "location": "人员简历表", "context": "[职称]"},
                {"name": "学历", "location": "人员简历表", "context": "[学历]"},
                {"name": "专业", "location": "人员简历表", "context": "[专业]"},
                {"name": "院校名称", "location": "人员简历表", "context": "[院校名称]"},
                {"name": "月数/年数", "location": "服务承诺", "context": "[月数/年数]"},
                {"name": "小时数", "location": "服务承诺", "context": "[小时数]"},
            ]

    # 为每个占位符添加类别
    for ph in placeholders:
        ph["category"] = classify_placeholder(ph["name"])

    return {
        "template_type": template_type,
        "template_name": name_map.get(template_type, f"{template_type}文件"),
        "file_path": file_path,
        "placeholders": placeholders,
        "project_info": {
            "id": project_info.get("id"),
            "title": project_info.get("title"),
            "project_name": project_info.get("project_name"),
            "bid_number": project_info.get("bid_number"),
        },
        "generated_at": datetime.now().isoformat(),
    }


def infer_project_type(parsed_data: Dict[str, Any]) -> str:
    """根据解析数据推断项目类型。"""
    full_text = json.dumps(parsed_data, ensure_ascii=False) if isinstance(parsed_data, dict) else str(parsed_data)

    best_type = "服务"  # 默认
    best_score = 0

    for proj_type, keywords in PROJECT_TYPE_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in full_text)
        if score > best_score:
            best_score = score
            best_type = proj_type

    return best_type


# ==================== 数据库操作 ====================

def save_templates_to_db(project_id: int, templates: List[Dict[str, Any]],
                         get_db_connection_func) -> List[Dict[str, Any]]:
    """将生成的模板保存到 bid_templates 表。"""
    saved = []
    with get_db_connection_func() as conn:
        cursor = conn.cursor()

        # 清理该项目之前的模板
        cursor.execute("DELETE FROM bid_templates WHERE content->'project_info'->>'id' = %s",
                       (str(project_id),))

        for tpl in templates:
            cursor.execute(
                """
                INSERT INTO bid_templates (template_type, template_name, content, is_default)
                VALUES (%s, %s, %s, %s)
                RETURNING id, template_type, template_name
                """,
                (
                    tpl["template_type"],
                    tpl["template_name"],
                    json.dumps(tpl, ensure_ascii=False),
                    True,
                ),
            )
            row = cursor.fetchone()
            saved.append({
                "id": row[0],
                "template_type": row[1],
                "template_name": row[2],
                "file_path": tpl.get("file_path"),
                "placeholder_count": len(tpl.get("placeholders", [])),
            })

        conn.commit()
        cursor.close()

    return saved


def load_templates_from_db(project_id: int, get_db_connection_func) -> List[Dict[str, Any]]:
    """从 bid_templates 表加载模板。"""
    with get_db_connection_func() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, template_type, template_name, content, created_at
            FROM bid_templates
            WHERE content->'project_info'->>'id' = %s
            ORDER BY template_type
            """,
            (str(project_id),),
        )
        rows = cursor.fetchall()
        cursor.close()

    results = []
    for row in rows:
        content = json.loads(row[3]) if isinstance(row[3], str) else row[3]
        results.append({
            "id": row[0],
            "template_type": row[1],
            "template_name": row[2],
            "content": content,
            "created_at": str(row[4]),
        })

    return results
