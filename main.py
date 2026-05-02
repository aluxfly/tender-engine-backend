"""
投标公司赚钱引擎 MVP 后端 API
FastAPI 单文件实现 - PostgreSQL 数据库

Day 1 审查修复:
  1. 连接池 (psycopg2.pool.ThreadedConnectionPool)
  2. 初始化断裂 (lifespan 中调用迁移脚本)
  3. UPDATE SQL 注入防护 (白名单)
  4. API Key 认证
  5. DELETE 资源泄漏修复 (try-finally)
  6. CORS 配置 (已有 *)
  7. 分页支持
  8. UNIQUE 约束 (迁移脚本)
  9. 统一错误响应
  10. doc_parser O(n²) 优化

Day 3 新功能:
  11. 文件上传模块 (file_uploader.py) — 多格式资料上传
  12. OCR 识别引擎 (ocr_engine.py) — PaddleOCR 营业执照/证书/通用识别
  13. 定时清理 (cleanup.py) — 24h 过期文件自动清理

Day 4 新功能:
  14. 标书整合引擎 (bid_merger.py) — 三套标书合并 + 封面 + 目录 + 页眉页脚
  15. PDF 导出 (pdf_exporter.py) — LibreOffice/pandoc/weasyprint 多策略转换
  16. 下载链接生成 — 24h 有效 token + bid_downloads 表
"""

from fastapi import FastAPI, HTTPException, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from contextlib import asynccontextmanager, contextmanager
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import random
from datetime import datetime
from docx import Document
from io import BytesIO
from fastapi.responses import StreamingResponse
import psycopg2
import psycopg2.extras
import psycopg2.pool
from pathlib import Path
import logging
import json
import os
import urllib.parse

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ==================== 数据库配置 ====================

# 全局连接池
db_pool: Optional[psycopg2.pool.ThreadedConnectionPool] = None


def get_database_url():
    """获取数据库连接 URL"""
    url = os.environ.get('DATABASE_URL')
    if url:
        return url
    return None


@contextmanager
def get_db_connection():
    """从连接池获取连接（上下文管理器，确保自动归还）"""
    global db_pool
    if db_pool is None:
        raise RuntimeError("数据库连接池未初始化")
    conn = db_pool.getconn()
    try:
        yield conn
    except Exception:
        conn.rollback()
        raise
    finally:
        db_pool.putconn(conn)


def verify_api_key(x_api_key: Optional[str] = Header(None)):
    """API Key 验证依赖"""
    api_key = os.environ.get('API_KEY', 'dev-key-2026')
    if not x_api_key or x_api_key != api_key:
        raise HTTPException(status_code=401, detail="无效的 API Key")


# ==================== 统一错误响应 ====================

def error_response(code: int, message: str, detail: str = "") -> dict:
    """统一错误响应格式"""
    return {"code": code, "message": message, "detail": detail}


# ==================== 应用生命周期 ====================

@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    global db_pool
    logger.info("应用启动中...")

    # 初始化连接池
    db_url = get_database_url()
    if db_url:
        min_conn = int(os.environ.get('DB_POOL_MIN', '5'))
        max_conn = int(os.environ.get('DB_POOL_MAX', '20'))
        db_pool = psycopg2.pool.ThreadedConnectionPool(
            minconn=min_conn, maxconn=max_conn, dsn=db_url
        )
        logger.info(f"连接池初始化完成 (min={min_conn}, max={max_conn})")
    else:
        logger.warning("DATABASE_URL 未设置，连接池未初始化")

    # 初始化数据库 + 迁移标书表
    init_database()

    # 启动定时清理任务
    try:
        from cleanup import start_cleanup_scheduler
        start_cleanup_scheduler()
        logger.info("定时清理任务已注册")
    except Exception as e:
        logger.warning(f"定时清理任务启动失败: {e}")

    logger.info("应用启动完成")
    yield
    # 关闭连接池
    if db_pool:
        db_pool.closeall()
        logger.info("连接池已关闭")
    # 停止定时清理
    try:
        from cleanup import stop_cleanup_scheduler
        stop_cleanup_scheduler()
    except Exception:
        pass
    logger.info("应用关闭")


app = FastAPI(
    title="投标公司赚钱引擎 API",
    description="MVP 版本 - 项目查询、中标预测、标书生成、资料上传、OCR识别",
    version="1.5.0",  # bump for Day 4: bid merger + PDF export + download links
    lifespan=lifespan
)

# CORS 配置（已存在，覆盖所有新增 API）
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 静态文件目录
STATIC_DIR = Path(__file__).parent / 'static'


# ==================== Day 3 路由注册 ====================

from file_uploader import router as material_router
from ocr_engine import router as ocr_router

app.include_router(material_router)
app.include_router(ocr_router)

logger.info("Day 3 路由已注册: /api/bid/material, /api/bid/ocr")


# ==================== 前端页面路由 ====================

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    """提供首页"""
    index_path = STATIC_DIR / 'index.html'
    if index_path.exists():
        with open(index_path, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())
    raise HTTPException(status_code=404, detail="Page not found")


@app.get("/projects", response_class=HTMLResponse)
async def serve_projects():
    """提供项目列表页"""
    page_path = STATIC_DIR / 'projects.html'
    if page_path.exists():
        with open(page_path, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())
    raise HTTPException(status_code=404, detail="Page not found")


@app.get("/project-detail", response_class=HTMLResponse)
async def serve_project_detail():
    """提供项目详情页"""
    page_path = STATIC_DIR / 'project-detail.html'
    if page_path.exists():
        with open(page_path, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())
    raise HTTPException(status_code=404, detail="Page not found")


@app.get("/bid", response_class=HTMLResponse)
async def serve_bid():
    """提供标书生成页"""
    page_path = STATIC_DIR / 'bid.html'
    if page_path.exists():
        with open(page_path, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())
    raise HTTPException(status_code=404, detail="Page not found")


def init_database():
    """初始化数据库 - 创建表、加载初始数据、执行迁移脚本"""
    if db_pool is None:
        logger.warning("DATABASE_URL 未设置，跳过数据库初始化")
        return

    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()

            # 检查 bid_notices 表是否存在
            cursor.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables
                    WHERE table_name = 'bid_notices'
                )
            """)
            table_exists = cursor.fetchone()[0]

            if not table_exists:
                logger.info("创建 bid_notices 表...")
                cursor.execute('''
                    CREATE TABLE bid_notices (
                        id SERIAL PRIMARY KEY,
                        title TEXT NOT NULL,
                        region TEXT,
                        budget REAL,
                        deadline TEXT,
                        description TEXT,
                        source_url TEXT,
                        source_site TEXT,
                        source TEXT,
                        category TEXT,
                        publish_date TEXT,
                        crawl_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        content_hash TEXT,
                        project_code TEXT,
                        status TEXT
                    )
                ''')
                conn.commit()
                logger.info("bid_notices 表创建成功")
            else:
                # 迁移：检查并添加缺失的列
                cursor.execute("""
                    SELECT column_name FROM information_schema.columns
                    WHERE table_name = 'bid_notices'
                """)
                columns = [row[0] for row in cursor.fetchall()]
                migration_cols = {
                    'source': 'TEXT',
                    'content_hash': 'TEXT',
                    'project_code': 'TEXT',
                    'status': 'TEXT'
                }
                for col_name, col_type in migration_cols.items():
                    if col_name not in columns:
                        logger.info(f"添加 {col_name} 字段...")
                        cursor.execute(f"ALTER TABLE bid_notices ADD COLUMN {col_name} {col_type}")
                        conn.commit()

            # 检查数据是否为空
            cursor.execute('SELECT COUNT(*) FROM bid_notices')
            count = cursor.fetchone()[0]

            if count == 0:
                logger.info("数据库为空，加载初始数据...")
                initial_data_path = Path(__file__).parent / 'initial_data.json'

                if initial_data_path.exists():
                    with open(initial_data_path, 'r', encoding='utf-8') as f:
                        initial_data = json.load(f)

                    for item in initial_data:
                        cursor.execute('''
                            INSERT INTO bid_notices
                            (title, region, budget, deadline, description, source_url, source_site, source, category, publish_date)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        ''', (
                            item.get('title', ''),
                            item.get('region', ''),
                            item.get('budget'),
                            item.get('deadline', ''),
                            item.get('description', ''),
                            item.get('source_url', ''),
                            item.get('source_site', ''),
                            item.get('source', ''),
                            item.get('category', ''),
                            item.get('publish_date', '')
                        ))

                    conn.commit()
                    logger.info(f"已加载 {len(initial_data)} 条初始数据")
                else:
                    logger.warning("初始数据文件不存在")
            else:
                logger.info(f"数据库已有 {count} 条数据，跳过初始化")

            cursor.close()
            logger.info("数据库初始化完成")

    except Exception as e:
        logger.error(f"数据库初始化失败：{e}")

    # 执行标书表迁移脚本（修复 #2：初始化断裂）
    try:
        from migrate_bid_tables import migrate as migrate_bid
        migrate_bid()
        logger.info("标书表迁移完成")
    except Exception as e:
        logger.warning(f"标书表迁移失败（可能已存在）：{e}")

    # Day 4 迁移：bid_downloads 表
    try:
        from migrate_day4 import migrate as migrate_day4
        migrate_day4()
        logger.info("Day 4 迁移完成 (bid_downloads 表)")
    except Exception as e:
        logger.warning(f"Day 4 迁移失败（可能已存在）：{e}")


# ==================== 请求/响应模型 ====================

class PredictRequest(BaseModel):
    project_id: int
    company_name: Optional[str] = "默认投标公司"
    bid_amount: Optional[float] = None


class PredictResponse(BaseModel):
    project_id: int
    project_name: str
    prediction: str  # 高/中/低
    confidence: float
    advice: str


class BidGenerateRequest(BaseModel):
    project_id: int
    company_name: str
    contact_person: str
    contact_phone: str
    bid_amount: float
    project_type: Optional[str] = "物联网卡"  # 物联网卡/布控球
    custom_fields: Optional[Dict[str, Any]] = None  # 自定义字段


# ==================== API 端点 ====================

@app.get("/")
def root():
    """API 首页"""
    db_stats = {"total": 0}
    if db_pool:
        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('SELECT COUNT(*) FROM bid_notices')
                row = cursor.fetchone()
                if row:
                    db_stats["total"] = row[0]
                cursor.close()
        except Exception as e:
            logger.warning(f"数据库查询失败: {e}")

    return {
        "message": "投标公司赚钱引擎 API - 真实数据版本",
        "version": "1.5.0",
        "database": {
            "total_notices": db_stats["total"],
            "status": "active" if db_stats["total"] > 0 else "empty"
        },
        "endpoints": {
            "projects": "GET /api/projects",
            "predict": "POST /api/predict",
            "bid_generate": "POST /api/bid/generate",
            "stats": "GET /api/stats",
            "crawl": "POST /api/crawl",
            "reload_data": "POST /api/reload-data"
        }
    }


@app.get("/api/projects")
def get_projects(category: Optional[str] = None, location: Optional[str] = None, limit: int = 50):
    """
    获取项目列表（真实招标数据）

    - category: 按类别筛选（可选）
    - location: 按地点筛选（可选）
    - limit: 返回数量限制（默认 50）
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            query = '''
                SELECT id, title as name, budget, deadline, category, region as location,
                       description, source_url, source_site, source, publish_date
                FROM bid_notices
                WHERE 1=1
            '''
            params = []

            if category:
                query += ' AND category LIKE %s'
                params.append(f'%{category}%')

            if location:
                query += ' AND region LIKE %s'
                params.append(f'%{location}%')

            query += ' ORDER BY deadline ASC LIMIT %s'
            params.append(limit)

            cursor.execute(query, params)
            rows = cursor.fetchall()

            results = []
            for row in rows:
                results.append({
                    'id': row['id'],
                    'name': row['name'],
                    'budget': row['budget'],
                    'deadline': row['deadline'],
                    'category': row['category'],
                    'location': row['location'],
                    'description': row['description'],
                    'source_url': row['source_url'],
                    'source_site': row['source_site'],
                    'source': row['source'],
                    'publish_date': row['publish_date']
                })

            return results

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取项目失败：{e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/projects/filter")
def filter_projects(keywords: Optional[str] = None, project_type: Optional[str] = None):
    """
    筛选项目（物联网卡/布控球）

    - keywords: 关键词，逗号分隔
    - project_type: 项目类型（物联网卡/布控球）
    """
    try:
        data_file = Path(__file__).parent / 'data.json'

        if not data_file.exists():
            return []

        with open(data_file, 'r', encoding='utf-8') as f:
            all_projects = json.load(f)

        if not keywords and not project_type:
            return all_projects[:15]

        filtered = []
        keyword_list = keywords.split(',') if keywords else []

        for project in all_projects:
            match = False

            if project_type and project.get('type') == project_type:
                match = True

            if keyword_list:
                project_name = project.get('name', '').lower()
                for kw in keyword_list:
                    if kw.lower() in project_name:
                        match = True
                        break

            if match:
                filtered.append(project)

        return filtered[:15]

    except Exception as e:
        logger.error(f"筛选项目失败：{e}")
        return []


@app.get("/api/bids")
def get_bids_alias(status: Optional[str] = None, limit: int = 50):
    """/api/bids 别名 → 转发到 /api/projects（前端兼容）"""
    return get_projects(category=None, location=None, limit=limit)


@app.post("/api/predict", response_model=PredictResponse)
def predict_bid_success(request: PredictRequest):
    """
    预测中标概率

    返回高/中/低三种预测结果
    """
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute('SELECT id, title FROM bid_notices WHERE id = %s', (request.project_id,))
            row = cursor.fetchone()
            cursor.close()

        if not row:
            raise HTTPException(status_code=404, detail=f"项目 ID {request.project_id} 不存在")

        project_name = row['title']
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询项目失败：{e}")
        raise HTTPException(status_code=500, detail=f"数据库查询失败：{str(e)}")

    predictions = [
        {"level": "高", "confidence": 0.85, "advice": "建议积极投标，中标概率较大"},
        {"level": "中", "confidence": 0.55, "advice": "可考虑投标，需优化报价策略"},
        {"level": "低", "confidence": 0.25, "advice": "谨慎评估，建议调整投标策略"}
    ]

    result = random.choice(predictions)

    return PredictResponse(
        project_id=request.project_id,
        project_name=project_name,
        prediction=result["level"],
        confidence=result["confidence"],
        advice=result["advice"]
    )


@app.post("/api/bid/generate", dependencies=[Depends(verify_api_key)])
def generate_bid_document(request: BidGenerateRequest):
    """
    生成投标文件（Word 格式）

    支持物联网卡和布控球两种模板
    返回可下载的 .docx 文件
    """
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                'SELECT id, title, budget, deadline, category, region, description FROM bid_notices WHERE id = %s',
                (request.project_id,))
            row = cursor.fetchone()
            cursor.close()

        if not row:
            raise HTTPException(status_code=404, detail=f"项目 ID {request.project_id} 不存在")

        project = {
            'id': row['id'],
            'name': row['title'],
            'budget': row['budget'],
            'deadline': row['deadline'],
            'category': row['category'],
            'location': row['region'],
            'description': row['description']
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询项目失败：{e}")
        raise HTTPException(status_code=500, detail=f"数据库查询失败：{str(e)}")

    template_type = request.project_type if request.project_type in ['物联网卡', '布控球'] else '物联网卡'

    template = None
    if template_type == '物联网卡':
        template_path = Path(__file__).parent / 'templates' / 'iot-sim-card-template.json'
    else:
        template_path = Path(__file__).parent / 'templates' / 'surveillance-ball-template.json'

    if template_path.exists():
        with open(template_path, 'r', encoding='utf-8') as f:
            template = json.load(f)

    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT

    def _set_font(run, name='宋体', size=12, bold=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        run._element.rPr.rFonts.set(qn('w:ascii'), name)

    def _add_para_with_style(doc, text='', style='Normal', alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              font_name='宋体', font_size=12, bold=False,
                              space_before=None, space_after=None, color=None,
                              first_line_indent=None):
        para = doc.add_paragraph(text, style=style)
        para.alignment = alignment
        for run in para.runs:
            _set_font(run, name=font_name, size=font_size, bold=bold, color=color)
        if space_before is not None:
            para.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
        if first_line_indent is not None:
            para.paragraph_format.first_line_indent = Cm(first_line_indent)
        return para

    def _add_styled_heading(doc, text, level=1, font_name='黑体', font_size=None,
                            space_before=12, space_after=6):
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            fs = font_size if font_size else ({1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12))
            _set_font(run, name=font_name, size=fs, bold=True)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        return heading

    def _add_kv_para(doc, key, value, key_width=4, font_size=11):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Cm(1)
        run_key = para.add_run(f'{key}：')
        _set_font(run_key, size=font_size, bold=True)
        run_val = para.add_run(str(value))
        _set_font(run_val, size=font_size)
        return para

    def _add_bullet(doc, text, level=0, font_size=11):
        para = doc.add_paragraph()
        indent = 1 + level * 1.5
        para.paragraph_format.left_indent = Cm(indent)
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        bullet = '•' if level == 0 else '◦'
        run_bullet = para.add_run(f'{bullet} ')
        _set_font(run_bullet, size=font_size, bold=(level == 0))
        run_text = para.add_run(text)
        _set_font(run_text, size=font_size)
        return para

    def _add_simple_table(doc, headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            _set_font(run, size=10, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = cell._element.get_or_add_tcPr()
            from lxml import etree
            shd = etree.SubElement(shading, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'D9E2F3')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                run = cell.paragraphs[0].add_run(str(val))
                _set_font(run, size=10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths:
            for ci, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[ci].width = Cm(w)
        return table

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    now_str = datetime.now().strftime('%Y年%m月%d日')

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(60)
    cover_title.paragraph_format.space_after = Pt(30)
    run_title = cover_title.add_run('投 标 文 件')
    _set_font(run_title, name='黑体', size=36, bold=True, color=(0, 51, 102))

    cover_subtitle = doc.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_subtitle.paragraph_format.space_after = Pt(40)
    run_sub = cover_subtitle.add_run(f'（{template_type}项目）')
    _set_font(run_sub, name='黑体', size=22, bold=False, color=(102, 102, 102))

    cover_proj = doc.add_paragraph()
    cover_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_proj.paragraph_format.space_after = Pt(50)
    run_p = cover_proj.add_run(f'项目名称：{project["name"]}')
    _set_font(run_p, name='楷体', size=16, bold=False)

    cover_comp = doc.add_paragraph()
    cover_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_comp.paragraph_format.space_after = Pt(50)
    run_c = cover_comp.add_run(f'投标单位：{request.company_name}')
    _set_font(run_c, name='楷体', size=16, bold=False)

    cover_date = doc.add_paragraph()
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_date.paragraph_format.space_after = Pt(20)
    run_d = cover_date.add_run(now_str)
    _set_font(run_d, name='楷体', size=14, bold=False)

    doc.add_page_break()

    # 目录页
    _add_styled_heading(doc, '目    录', level=1, font_size=18)
    doc.add_paragraph()
    toc_items = [
        ('第一章', '投标函'),
        ('第二章', '法定代表人授权书'),
        ('第三章', '投标报价一览表'),
        ('第四章', '技术方案'),
        ('第五章', '项目实施计划'),
        ('第六章', '售后服务方案'),
        ('第七章', '企业资质与业绩'),
        ('第八章', '项目理解与需求分析'),
    ]
    for num, title_text in toc_items:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_before = Pt(4)
        toc_para.paragraph_format.space_after = Pt(4)
        run_num = toc_para.add_run(f'{num}  ')
        _set_font(run_num, size=12, bold=True)
        run_text = toc_para.add_run(title_text)
        _set_font(run_text, size=12)
    doc.add_page_break()

    # 第一章：投标函
    _add_styled_heading(doc, '第一章  投标函', level=1)
    doc.add_paragraph()

    _add_para_with_style(doc, f'致：{project["name"]}招标方', font_size=12, bold=False,
                          space_before=12, space_after=6, first_line_indent=0)
    _add_para_with_style(doc,
        f'    根据贵方发布的招标文件，{request.company_name}（以下简称"我方"）经认真研究招标文件的全部内容后，' +
        f'决定参与本项目的投标。现正式提交投标文件，具体内容如下：',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_para_with_style(doc,
        f'    一、我方完全理解并接受招标文件的全部内容和要求，愿意按照招标文件的规定提供{template_type}相关产品与服务。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_para_with_style(doc,
        f'    二、我方投标总报价为人民币（大写）：¥{request.bid_amount:,.2f}（含税）。该报价包含设备费、运输费、安装调试费、培训费、' +
        f'税费及售后服务等全部费用。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_para_with_style(doc,
        f'    三、我方承诺在中标后按照招标文件及合同约定履行全部责任和义务，保证按期、按质完成项目交付。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_para_with_style(doc,
        f'    四、我方投标文件自投标截止日起有效期为 90 个日历日。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_para_with_style(doc,
        f'    五、如我方中标，我方承诺在合同签订后按照约定时间完成项目交付，并提供完善的售后服务。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    doc.add_paragraph()
    _add_para_with_style(doc, f'投标单位（盖章）：{request.company_name}', font_size=12, first_line_indent=0, space_before=12, space_after=6)
    _add_para_with_style(doc, f'法定代表人或授权代表（签字）：{request.contact_person}', font_size=12, first_line_indent=0, space_before=6, space_after=6)
    _add_para_with_style(doc, f'联系电话：{request.contact_phone}', font_size=12, first_line_indent=0, space_before=6, space_after=6)
    _add_para_with_style(doc, f'日期：{now_str}', font_size=12, first_line_indent=0, space_before=6, space_after=6)

    # 第二章：法定代表人授权书
    _add_styled_heading(doc, '第二章  法定代表人授权书', level=1)
    doc.add_paragraph()

    _add_para_with_style(doc,
        f'    本授权书声明：{request.company_name}的法定代表人（姓名、职务）在此授权{request.contact_person}（姓名、职务）' +
        f'为我方合法代理人，以我方名义参加"{project["name"]}"项目的投标活动。' +
        f'代理人在投标过程中签署的一切文件和处理与之有关的一切事务，我方均予以承认。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    doc.add_paragraph()
    _add_kv_para(doc, '授权单位', request.company_name)
    _add_kv_para(doc, '被授权人', request.contact_person)
    _add_kv_para(doc, '联系电话', request.contact_phone)
    _add_kv_para(doc, '项目名称', project['name'])
    _add_kv_para(doc, '项目地点', project['location'])
    _add_kv_para(doc, '授权日期', now_str)

    doc.add_paragraph()
    _add_para_with_style(doc, '法定代表人签字：_____________    被授权人签字：_____________', font_size=12, first_line_indent=0, space_before=12, space_after=6)

    doc.add_page_break()

    # 第三章：投标报价一览表
    _add_styled_heading(doc, '第三章  投标报价一览表', level=1)
    doc.add_paragraph()

    _add_kv_para(doc, '项目名称', project['name'])
    _add_kv_para(doc, '项目编号', f'BID-{project["id"]}-{datetime.now().strftime("%Y%m%d")}')
    _add_kv_para(doc, '项目预算', f'¥{project["budget"]:,.2f}')
    _add_kv_para(doc, '投标报价', f'¥{request.bid_amount:,.2f}')

    budget_ratio = (request.bid_amount / project['budget'] * 100) if project['budget'] and project['budget'] > 0 else 0
    _add_kv_para(doc, '报价占预算比例', f'{budget_ratio:.1f}%')

    doc.add_paragraph()
    _add_styled_heading(doc, '报价明细表', level=3, font_size=13)
    doc.add_paragraph()

    if template_type == '物联网卡':
        cf = request.custom_fields or {}
        qty = int(cf.get('delivery_quantity', 1000))
        unit_price = request.bid_amount / qty if qty > 0 else request.bid_amount
        headers = ['序号', '费用项目', '单价（元）', '数量', '金额（元）', '备注']
        rows = [
            ['1', '物联网卡（SIM卡）', f'{unit_price * 0.4:.2f}', f'{qty}', f'{request.bid_amount * 0.4:.2f}', cf.get('card_type', '标准SIM卡')],
            ['2', '流量套餐费', f'{unit_price * 0.25:.2f}', f'{qty}', f'{request.bid_amount * 0.25:.2f}', cf.get('data_plan', '月流量1GB')],
            ['3', '平台接入费', f'{request.bid_amount * 0.1:.2f}', '1', f'{request.bid_amount * 0.1:.2f}', '卡管理平台'],
            ['4', 'API对接费', f'{request.bid_amount * 0.05:.2f}', '1', f'{request.bid_amount * 0.05:.2f}', '接口开发'],
            ['5', '技术支持服务费', f'{request.bid_amount * 0.1:.2f}', '1', f'{request.bid_amount * 0.1:.2f}', '含售后'],
            ['6', '运输及杂费', f'{request.bid_amount * 0.05:.2f}', '1', f'{request.bid_amount * 0.05:.2f}', '物流配送'],
            ['合计', '', '', '', f'{request.bid_amount:.2f}', '含税'],
        ]
    else:
        cf = request.custom_fields or {}
        qty = int(cf.get('delivery_quantity', 50))
        unit_price = request.bid_amount / qty if qty > 0 else request.bid_amount
        headers = ['序号', '费用项目', '单价（元）', '数量', '金额（元）', '备注']
        rows = [
            ['1', '布控球设备', f'{unit_price * 0.5:.2f}', f'{qty}', f'{request.bid_amount * 0.5:.2f}', cf.get('video_resolution', '4K超清')],
            ['2', '安装材料及配件', f'{unit_price * 0.1:.2f}', f'{qty}', f'{request.bid_amount * 0.1:.2f}', '支架、线缆等'],
            ['3', '云存储服务费', f'{request.bid_amount * 0.1:.2f}', '1', f'{request.bid_amount * 0.1:.2f}', cf.get('storage', '云存储')],
            ['4', '安装调试费', f'{request.bid_amount * 0.1:.2f}', '1', f'{request.bid_amount * 0.1:.2f}', '现场实施'],
            ['5', '智能分析模块', f'{request.bid_amount * 0.1:.2f}', '1', f'{request.bid_amount * 0.1:.2f}', 'AI算法'],
            ['6', '培训及运维费', f'{request.bid_amount * 0.05:.2f}', '1', f'{request.bid_amount * 0.05:.2f}', '培训+运维'],
            ['合计', '', '', '', f'{request.bid_amount:.2f}', '含税'],
        ]

    _add_simple_table(doc, headers, rows, col_widths=[1.5, 3.5, 2.5, 2, 2.5, 3])

    doc.add_paragraph()
    _add_para_with_style(doc,
        '注：以上报价为含税总价，报价有效期为 90 个日历日，自投标截止日起计算。',
        font_size=10, first_line_indent=0, space_before=6, space_after=6)

    doc.add_page_break()

    # 第四章：技术方案
    _add_styled_heading(doc, '第四章  技术方案', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '4.1 项目概述', level=2, font_size=14)
    _add_para_with_style(doc,
        f'    本项目为"{project["name"]}"，项目地点位于{project["location"]}，' +
        f'项目预算为¥{project["budget"]:,.2f}。{project.get("description", "")}。' +
        f'我方凭借在该领域的丰富经验和专业技术实力，提供完整的技术解决方案。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_styled_heading(doc, '4.2 技术路线与架构', level=2, font_size=14)

    if template_type == '物联网卡':
        _add_para_with_style(doc,
            '    我方采用"终端-网络-平台"三层架构的物联网卡解决方案，确保系统稳定可靠：',
            font_size=12, first_line_indent=0, space_before=6, space_after=6)
        _add_bullet(doc, '终端层：支持多类型物联网卡（SIM/eSIM/M2M），适配各类物联网设备')
        _add_bullet(doc, '网络层：支持 FDD-LTE/TDD-LTE/NB-IoT/5G 全频段覆盖，保障网络质量')
        _add_bullet(doc, '平台层：自建卡管理平台，提供卡生命周期管理、用量监控、故障诊断等核心功能')
        _add_bullet(doc, '安全层：端到端加密传输，双向认证机制，确保数据安全')
    else:
        _add_para_with_style(doc,
            '    我方采用"前端采集-网络传输-云端分析"三层架构的布控球解决方案：',
            font_size=12, first_line_indent=0, space_before=6, space_after=6)
        _add_bullet(doc, '前端采集层：4K超高清布控球设备，支持多种供电方式（太阳能/市电/电池）')
        _add_bullet(doc, '网络传输层：4G/5G无线传输，支持RTSP/ONVIF/GB/T28181标准协议')
        _add_bullet(doc, '云端分析层：视频云平台+AI智能分析，支持人脸识别、车牌识别、行为分析')
        _add_bullet(doc, '安全层：视频流加密传输，权限分级管控，符合等保要求')

    _add_styled_heading(doc, '4.3 核心技术参数', level=2, font_size=14)
    doc.add_paragraph()

    if template_type == '物联网卡':
        cf = request.custom_fields or {}
        spec_headers = ['参数项', '技术指标', '说明']
        spec_rows = [
            ['卡类型', cf.get('card_type', 'SIM卡（可定制eSIM/M2M）'), '支持多种卡类型'],
            ['频段支持', cf.get('frequency_band', '全频段（FDD-LTE/TDD-LTE/NB-IoT/5G）'), '广泛兼容性'],
            ['流量套餐', cf.get('data_plan', '月流量1GB（可定制）'), '灵活套餐选择'],
            ['运营商合作', cf.get('operator', '中国移动/联通/电信'), '多运营商冗余'],
            ['APN配置', cf.get('apn', '支持自定义APN'), '灵活网络配置'],
            ['QoS保障', cf.get('qos', '优先服务级别'), '网络质量保障'],
            ['安全加密', cf.get('security', '双向认证+端到端加密'), '高安全等级'],
            ['VPLMN配置', cf.get('vplmn', '支持46001/46003等'), '漫游支持'],
        ]
    else:
        cf = request.custom_fields or {}
        spec_headers = ['参数项', '技术指标', '说明']
        spec_rows = [
            ['视频分辨率', cf.get('video_resolution', '4K超高清（3840×2160）'), '高清画质'],
            ['帧率', cf.get('frame_rate', '30fps（支持60fps）'), '流畅视频'],
            ['码率', cf.get('bit_rate', '4Mbps-10Mbps自适应'), '智能码率控制'],
            ['存储方式', cf.get('storage', '本地+云存储双模式'), '可靠存储'],
            ['供电方式', cf.get('power', '太阳能+电池混合供电'), '无市电场景适用'],
            ['安装方式', cf.get('installation', '立杆/墙面/车载多模式'), '灵活部署'],
            ['无线传输', cf.get('wireless', '4G/5G/Wi-Fi混合组网'), '多网络冗余'],
            ['协议支持', cf.get('protocol', 'RTSP/ONVIF/GB/T28181'), '标准协议兼容'],
            ['智能分析', '人脸识别/车牌识别/行为分析', 'AI算法支持'],
        ]

    _add_simple_table(doc, spec_headers, spec_rows, col_widths=[3.5, 6, 4])

    _add_styled_heading(doc, '4.4 平台功能', level=2, font_size=14)

    if template_type == '物联网卡':
        _add_bullet(doc, '卡管理：开卡、停卡、换卡、销卡全生命周期管理')
        _add_bullet(doc, '用量监控：实时监控流量使用情况，预警提醒')
        _add_bullet(doc, 'API接口：提供RESTful API，支持第三方系统对接')
        _add_bullet(doc, '报表统计：多维度数据分析，支持自定义报表')
        _add_bullet(doc, '故障诊断：自动故障检测与告警，快速定位问题')
    else:
        _add_bullet(doc, '视频监控：实时画面查看、录像回放、截图取证')
        _add_bullet(doc, 'AI分析：人脸识别、车牌识别、行为分析等智能算法')
        _add_bullet(doc, '设备管理：设备在线状态监控、远程配置、OTA升级')
        _add_bullet(doc, '告警通知：智能告警规则设置，多渠道通知（短信/邮件/APP）')
        _add_bullet(doc, '数据大屏：可视化大屏展示，实时数据监控')

    doc.add_page_break()

    # 第五章：项目实施计划
    _add_styled_heading(doc, '第五章  项目实施计划', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '5.1 实施流程', level=2, font_size=14)
    _add_para_with_style(doc,
        '    我方将按照以下阶段推进项目实施：',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    implementation_headers = ['阶段', '时间', '主要工作内容', '交付物']
    if template_type == '物联网卡':
        implementation_rows = [
            ['第一阶段\n需求确认', '第1周', '需求调研、方案确认、APN配置确认', '需求确认书'],
            ['第二阶段\n卡片生产', '第2-3周', '卡片定制生产、质量检测、编号管理', '产品检测报告'],
            ['第三阶段\n平台部署', '第3-4周', '管理平台部署、API对接调试', '平台部署报告'],
            ['第四阶段\n物流交付', '第4-5周', '产品包装、物流配送、到货签收', '签收单'],
            ['第五阶段\n联调验收', '第5-6周', '现场联调测试、性能验证、项目验收', '验收报告'],
        ]
    else:
        implementation_rows = [
            ['第一阶段\n需求调研', '第1周', '现场勘测、需求确认、方案设计', '需求调研报告'],
            ['第二阶段\n设备采购', '第2-3周', '设备采购、出厂检测、配件准备', '设备检测报告'],
            ['第三阶段\n安装部署', '第3-5周', '现场安装、网络配置、平台对接', '安装部署报告'],
            ['第四阶段\n系统联调', '第5-6周', '系统联调、功能测试、性能优化', '测试报告'],
            ['第五阶段\n培训验收', '第6-7周', '用户培训、项目验收、交付文档', '验收报告'],
        ]

    _add_simple_table(doc, implementation_headers, implementation_rows, col_widths=[3, 2, 5.5, 3])

    _add_styled_heading(doc, '5.2 项目团队', level=2, font_size=14)
    _add_para_with_style(doc,
        '    我方将组建专业项目团队，确保项目高质量交付：',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    team_headers = ['角色', '人数', '职责']
    team_rows = [
        ['项目经理', '1人', '项目统筹管理、进度控制、对外沟通'],
        ['技术负责人', '1人', '技术方案设计、技术难题攻关'],
        ['实施工程师', '2-4人', '现场安装部署、系统调试'],
        ['测试工程师', '1人', '功能测试、性能测试、验收测试'],
        ['售后工程师', '1人', '售后支持、运维保障'],
    ]
    _add_simple_table(doc, team_headers, team_rows, col_widths=[3, 2, 8.5])

    _add_styled_heading(doc, '5.3 质量保障措施', level=2, font_size=14)
    _add_bullet(doc, '严格执行 ISO 9001 质量管理体系')
    _add_bullet(doc, '每阶段设立质量检查点，不合格返工')
    _add_bullet(doc, '关键节点邀请第三方检测机构参与')
    _add_bullet(doc, '建立项目日报/周报机制，及时沟通项目进展')

    doc.add_page_break()

    # 第六章：售后服务方案
    _add_styled_heading(doc, '第六章  售后服务方案', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '6.1 服务承诺', level=2, font_size=14)
    _add_bullet(doc, '质保期：自验收合格之日起 12 个月')
    _add_bullet(doc, '响应时间：7×24小时响应，2小时内远程响应，24小时内现场到达')
    _add_bullet(doc, '故障处理：一般故障 4 小时内解决，重大故障 24 小时内提供备用方案')
    _add_bullet(doc, '定期巡检：每季度一次现场巡检，提交巡检报告')
    _add_bullet(doc, '技术支持：免费技术培训 2 次/年，不定期技术交流')

    _add_styled_heading(doc, '6.2 售后服务体系', level=2, font_size=14)
    _add_para_with_style(doc,
        '    我方建立了完善的三级售后服务体系：',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)
    _add_bullet(doc, '一级支持：客服热线、在线工单系统、自助知识库')
    _add_bullet(doc, '二级支持：远程诊断、远程修复、远程升级')
    _add_bullet(doc, '三级支持：现场服务、备件更换、设备维修')

    _add_styled_heading(doc, '6.3 备品备件保障', level=2, font_size=14)
    _add_para_with_style(doc,
        '    我方承诺在质保期内免费提供备品备件服务。对于关键设备，' +
        '我方将在项目所在地设立备件库，确保故障设备能快速替换，不影响项目运行。',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_styled_heading(doc, '6.4 培训计划', level=2, font_size=14)
    training_headers = ['培训阶段', '培训内容', '培训时长', '培训对象']
    training_rows = [
        ['第一阶段', '产品功能与操作', '4学时', '操作人员'],
        ['第二阶段', '日常维护与故障排查', '4学时', '运维人员'],
        ['第三阶段', '高级功能与系统管理', '4学时', '管理人员'],
    ]
    _add_simple_table(doc, training_headers, training_rows, col_widths=[3, 5, 2.5, 3])

    doc.add_page_break()

    # 第七章：企业资质与业绩
    _add_styled_heading(doc, '第七章  企业资质与业绩', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '7.1 企业资质', level=2, font_size=14)

    if template_type == '物联网卡':
        qual_headers = ['资质名称', '资质说明']
        qual_rows = [
            ['营业执照', '合法有效，经营范围包含物联网相关业务'],
            ['增值电信业务经营许可证', '含ICP许可证、EDI证书'],
            ['物联网相关资质', '物联网行业准入资质'],
            ['ISO 9001 质量管理体系认证', '通过国际质量体系认证'],
            ['ISO 27001 信息安全管理体系', '通过信息安全管理体系认证'],
        ]
    else:
        qual_headers = ['资质名称', '资质说明']
        qual_rows = [
            ['营业执照', '合法有效，经营范围包含安防相关业务'],
            ['安防工程企业资质', '含安防设计施工维护资质'],
            ['通信工程施工资质', '通信设备安装调试资质'],
            ['信息系统集成资质', '系统集成能力认证'],
            ['ISO 9001 质量管理体系认证', '通过国际质量体系认证'],
        ]
    _add_simple_table(doc, qual_headers, qual_rows, col_widths=[5, 8.5])

    _add_styled_heading(doc, '7.2 类似项目业绩', level=2, font_size=14)
    _add_para_with_style(doc,
        f'    {request.company_name}在{template_type}领域积累了丰富的项目经验，以下为部分代表性案例：',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    if template_type == '物联网卡':
        case_headers = ['项目名称', '项目规模', '服务内容', '完成时间']
        case_rows = [
            ['某省电力物联网卡项目', '50000张', '物联网卡供应+平台接入', '2025年'],
            ['某市智慧交通物联网卡项目', '30000张', '物联网卡+API对接', '2025年'],
            ['某物流企业物联网卡项目', '20000张', '物联网卡+管理平台', '2024年'],
            ['某农业物联网监测项目', '10000张', 'NB-IoT卡+云平台', '2024年'],
        ]
    else:
        case_headers = ['项目名称', '项目规模', '服务内容', '完成时间']
        case_rows = [
            ['某市雪亮工程布控球项目', '200台', '设备供应+安装+平台', '2025年'],
            ['某景区智能监控布控球项目', '80台', '设备+AI分析+云存储', '2025年'],
            ['某工地安全监控布控球项目', '150台', '设备供应+安装运维', '2024年'],
            ['某河道防洪监控布控球项目', '60台', '太阳能布控球+云平台', '2024年'],
        ]
    _add_simple_table(doc, case_headers, case_rows, col_widths=[4.5, 2.5, 4.5, 2])

    doc.add_page_break()

    # 第八章：项目理解与需求分析
    _add_styled_heading(doc, '第八章  项目理解与需求分析', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '8.1 项目背景理解', level=2, font_size=14)
    _add_para_with_style(doc,
        f'    通过对"{project["name"]}"项目的深入研究，我方理解该项目的主要背景为：',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    {project.get("description", "本项目旨在通过采购相关设备与服务，提升业务能力和管理效率。")}',
        font_size=12, first_line_indent=0, space_before=6, space_after=6)

    _add_styled_heading(doc, '8.2 核心需求分析', level=2, font_size=14)

    if template_type == '物联网卡':
        _add_bullet(doc, '稳定可靠的网络连接：保障物联网设备全天候在线')
        _add_bullet(doc, '灵活的流量管理：根据业务需求提供差异化流量方案')
        _add_bullet(doc, '完善的平台支撑：实现卡的集中管理和远程运维')
        _add_bullet(doc, '安全保障机制：确保数据传输和存储安全')
        _add_bullet(doc, '快速响应服务：建立高效的故障处理机制')
    else:
        _add_bullet(doc, '高质量视频采集：确保监控画面清晰流畅')
        _add_bullet(doc, '可靠的传输网络：保障视频数据实时传输')
        _add_bullet(doc, '智能分析能力：利用AI技术提升监控效率')
        _add_bullet(doc, '灵活的部署方案：适应不同场景的安装需求')
        _add_bullet(doc, '完善的售后服务：保障系统长期稳定运行')

    _add_styled_heading(doc, '8.3 我方优势', level=2, font_size=14)
    _add_bullet(doc, '行业经验丰富：在相关领域拥有多年项目经验')
    _add_bullet(doc, '技术实力雄厚：自有研发团队，持续技术创新')
    _add_bullet(doc, '服务网络完善：覆盖全国的服务网点')
    _add_bullet(doc, '成本优势明显：规模采购带来的价格优势')
    _add_bullet(doc, '质量保证体系：通过多项国际认证')

    # 附录
    doc.add_page_break()
    _add_styled_heading(doc, '附    录', level=1)
    doc.add_paragraph()
    _add_kv_para(doc, '文件生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    _add_kv_para(doc, '生成系统', '投标公司赚钱引擎 AI 标书生成系统 v2.0')
    _add_kv_para(doc, '项目类型', template_type)
    _add_kv_para(doc, '文档版本', 'V1.0')
    _add_kv_para(doc, '文档说明', '本文档由系统自动生成，仅供参考，正式投标前请由专业人员审核完善。')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{template_type}_bid_{project['id']}_{request.company_name}.docx"
    encoded_filename = urllib.parse.quote(filename, safe='')

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@app.get("/api/stats")
def get_stats():
    """获取数据库统计信息"""
    if not db_pool:
        return {"status": "error", "code": 503, "message": "数据库未连接", "detail": "DATABASE_URL 未设置"}
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            cursor.execute('SELECT COUNT(*) as count FROM bid_notices')
            total = cursor.fetchone()["count"]

            cursor.execute('SELECT source_site, COUNT(*) as count FROM bid_notices GROUP BY source_site')
            by_source = {row["source_site"]: row["count"] for row in cursor.fetchall()}

            cursor.execute('SELECT region, COUNT(*) as count FROM bid_notices WHERE region != \'\' GROUP BY region ORDER BY count DESC LIMIT 10')
            by_region = {row["region"]: row["count"] for row in cursor.fetchall()}

            cursor.execute('SELECT MAX(crawl_time) as last_update FROM bid_notices')
            last_update = cursor.fetchone()["last_update"]

            cursor.close()

        return {
            "total": total,
            "by_source": by_source,
            "by_region": by_region,
            "last_update": str(last_update) if last_update else None,
            "status": "success"
        }
    except Exception as e:
        logger.error(f"获取统计失败：{e}")
        return {"status": "error", "code": 500, "message": "获取统计失败", "detail": str(e)}


@app.post("/api/crawl", dependencies=[Depends(verify_api_key)])
def trigger_crawl():
    """手动触发爬虫任务"""
    try:
        from crawler.gov_crawler import crawl_all

        logger.info("手动触发爬虫任务...")
        results = crawl_all()

        return {
            "status": "success",
            "message": "爬虫任务执行完成",
            "results": results
        }
    except Exception as e:
        logger.error(f"爬虫任务失败：{e}")
        raise HTTPException(status_code=500, detail=f"爬虫执行失败：{str(e)}")


@app.post("/api/reload-data", dependencies=[Depends(verify_api_key)])
def reload_initial_data():
    """手动重新加载初始数据"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()

            cursor.execute("DELETE FROM bid_notices")
            conn.commit()
            logger.info("已清空现有数据...")

            initial_data_path = Path(__file__).parent / 'initial_data.json'

            if not initial_data_path.exists():
                raise HTTPException(status_code=500, detail="初始数据文件不存在")

            with open(initial_data_path, 'r', encoding='utf-8') as f:
                initial_data = json.load(f)

            for item in initial_data:
                cursor.execute('''
                    INSERT INTO bid_notices
                    (title, region, budget, deadline, description, source_url, source_site, source, category, publish_date)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ''', (
                    item.get('title', ''),
                    item.get('region', ''),
                    item.get('budget'),
                    item.get('deadline', ''),
                    item.get('description', ''),
                    item.get('source_url', ''),
                    item.get('source_site', ''),
                    item.get('source', ''),
                    item.get('category', ''),
                    item.get('publish_date', '')
                ))

            conn.commit()
            cursor.close()

        logger.info(f"已重新加载 {len(initial_data)} 条初始数据")

        return {
            "status": "success",
            "message": f"已重新加载 {len(initial_data)} 条初始数据",
            "count": len(initial_data)
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"重新加载数据失败：{e}")
        raise HTTPException(status_code=500, detail=f"重新加载失败：{str(e)}")


@app.post("/api/projects/bulk-import")
def bulk_import_projects(data: List[Dict[str, Any]]):
    """批量导入项目数据"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()

            cursor.execute("""
                SELECT column_name FROM information_schema.columns
                WHERE table_name = 'bid_notices'
            """)
            columns = [row[0] for row in cursor.fetchall()]
            alter_cols = ['project_code', 'status', 'source']
            for col in alter_cols:
                if col not in columns:
                    cursor.execute(f"ALTER TABLE bid_notices ADD COLUMN {col} TEXT")
                    conn.commit()

            imported = 0
            skipped = 0

            for item in data:
                code = item.get('project_code', '')
                if code:
                    cursor.execute("SELECT id FROM bid_notices WHERE project_code = %s", (code,))
                    if cursor.fetchone():
                        skipped += 1
                        continue

                title = item.get('title', '')
                publish_date = item.get('publish_date', '')
                cursor.execute(
                    "SELECT id FROM bid_notices WHERE title = %s AND publish_date = %s",
                    (title, publish_date))
                if cursor.fetchone():
                    skipped += 1
                    continue

                cursor.execute('''
                    INSERT INTO bid_notices
                    (title, region, budget, deadline, description, source_url, source_site, source, category, publish_date, project_code, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ''', (
                    title,
                    item.get('region', ''),
                    item.get('budget', 0),
                    item.get('deadline', ''),
                    item.get('description', ''),
                    item.get('source_url', ''),
                    item.get('source_site', ''),
                    item.get('source', ''),
                    item.get('category', ''),
                    publish_date,
                    code,
                    item.get('status', '')
                ))
                imported += 1

            conn.commit()
            cursor.close()

        return {
            "status": "success",
            "imported": imported,
            "skipped": skipped,
            "message": f"导入 {imported} 条，跳过 {skipped} 条（已存在）"
        }
    except Exception as e:
        logger.error(f"批量导入失败：{e}")
        raise HTTPException(status_code=500, detail=f"导入失败：{str(e)}")


# =====================================================================
# 标书 AI 生成系统 — Day 1 新增 API（已修复审查问题）
# =====================================================================

# ---- Pydantic 模型 ----

class CompanyProfileCreate(BaseModel):
    company_name: str
    credit_code: Optional[str] = None
    legal_representative: Optional[str] = None
    contact_person: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    address: Optional[str] = None
    bank_info: Optional[Dict[str, Any]] = None
    qualifications: Optional[List[Any]] = None


class CompanyProfileUpdate(BaseModel):
    company_name: Optional[str] = None
    credit_code: Optional[str] = None
    legal_representative: Optional[str] = None
    contact_person: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    address: Optional[str] = None
    bank_info: Optional[Dict[str, Any]] = None
    qualifications: Optional[List[Any]] = None


class BidProjectCreate(BaseModel):
    title: str
    source_file_name: Optional[str] = None
    file_path: Optional[str] = None
    parsed_data: Optional[Dict[str, Any]] = None
    status: Optional[str] = "draft"


# ---- 公司资料管理 API（修复 #4：添加 API Key 认证）----

@app.post("/api/company/profile", response_model=dict, dependencies=[Depends(verify_api_key)])
def create_company_profile(profile: CompanyProfileCreate):
    """创建公司资料（需要 API Key）"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            cursor.execute(
                """
                INSERT INTO company_profiles
                    (company_name, credit_code, legal_representative, contact_person,
                     phone, email, address, bank_info, qualifications)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id, company_name, credit_code, legal_representative,
                          contact_person, phone, email, address, bank_info, qualifications,
                          created_at, updated_at
                """,
                (
                    profile.company_name,
                    profile.credit_code,
                    profile.legal_representative,
                    profile.contact_person,
                    profile.phone,
                    profile.email,
                    profile.address,
                    json.dumps(profile.bank_info) if profile.bank_info else None,
                    json.dumps(profile.qualifications) if profile.qualifications else None,
                ),
            )
            row = cursor.fetchone()
            conn.commit()
            cursor.close()

        return {"status": "success", "data": dict(row)}

    except psycopg2.errors.UniqueViolation:
        raise HTTPException(status_code=409, detail="该公司资料已存在（信用代码重复）")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建公司资料失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/company/profile", response_model=dict, dependencies=[Depends(verify_api_key)])
def get_company_profiles(name: Optional[str] = None, limit: int = 20):
    """查询公司资料列表（需要 API Key）"""
    if not db_pool:
        return {"status": "error", "code": 503, "message": "数据库未连接", "detail": "DATABASE_URL 未设置"}
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            if name:
                cursor.execute(
                    """
                    SELECT id, company_name, credit_code, legal_representative,
                           contact_person, phone, email, address, bank_info, qualifications,
                           created_at, updated_at
                    FROM company_profiles
                    WHERE company_name LIKE %s
                    ORDER BY created_at DESC
                    LIMIT %s
                    """,
                    (f"%{name}%", limit),
                )
            else:
                cursor.execute(
                    """
                    SELECT id, company_name, credit_code, legal_representative,
                           contact_person, phone, email, address, bank_info, qualifications,
                           created_at, updated_at
                    FROM company_profiles
                    ORDER BY created_at DESC
                    LIMIT %s
                    """,
                    (limit,),
                )

            rows = cursor.fetchall()
            cursor.close()

        return {
            "status": "success",
            "count": len(rows),
            "data": [dict(r) for r in rows],
        }

    except Exception as e:
        logger.error(f"查询公司资料失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/company/profile/{profile_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def update_company_profile(profile_id: int, profile: CompanyProfileUpdate):
    """
    更新公司资料（需要 API Key）
    修复 #3：使用白名单过滤字段名，防止 SQL 注入
    修复 #5：使用 try-finally 确保连接正确归还
    """
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            # 检查是否存在
            cursor.execute("SELECT id FROM company_profiles WHERE id = %s", (profile_id,))
            if not cursor.fetchone():
                cursor.close()
                raise HTTPException(status_code=404, detail=f"公司资料 ID {profile_id} 不存在")

            # 修复 #3：白名单过滤字段名，防止 SQL 注入
            ALLOWED_UPDATE_FIELDS = {
                "company_name", "credit_code", "legal_representative",
                "contact_person", "phone", "email", "address"
            }
            updates = []
            values = []
            field_map = {
                "company_name": profile.company_name,
                "credit_code": profile.credit_code,
                "legal_representative": profile.legal_representative,
                "contact_person": profile.contact_person,
                "phone": profile.phone,
                "email": profile.email,
                "address": profile.address,
            }
            for field, val in field_map.items():
                if field not in ALLOWED_UPDATE_FIELDS:
                    continue  # 白名单外字段直接跳过
                if val is not None:
                    updates.append(f"{field} = %s")
                    values.append(val)

            if profile.bank_info is not None:
                updates.append("bank_info = %s")
                values.append(json.dumps(profile.bank_info))
            if profile.qualifications is not None:
                updates.append("qualifications = %s")
                values.append(json.dumps(profile.qualifications))

            if not updates:
                cursor.close()
                raise HTTPException(status_code=400, detail="没有提供要更新的字段")

            updates.append("updated_at = CURRENT_TIMESTAMP")
            values.append(profile_id)

            cursor.execute(
                f"UPDATE company_profiles SET {', '.join(updates)} WHERE id = %s "
                "RETURNING id, company_name, credit_code, legal_representative, "
                "contact_person, phone, email, address, bank_info, qualifications, "
                "created_at, updated_at",
                values,
            )
            row = cursor.fetchone()
            conn.commit()
            cursor.close()

        return {"status": "success", "data": dict(row)}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"更新公司资料失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/company/profile/{profile_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def delete_company_profile(profile_id: int):
    """
    删除公司资料（需要 API Key）
    修复 #5：使用上下文管理器确保连接总是被归还
    """
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM company_profiles WHERE id = %s RETURNING id", (profile_id,))
            row = cursor.fetchone()
            conn.commit()
            cursor.close()

        if not row:
            raise HTTPException(status_code=404, detail=f"公司资料 ID {profile_id} 不存在")

        return {"status": "success", "message": f"已删除公司资料 ID {profile_id}"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除公司资料失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ---- 标书项目 API ----

@app.post("/api/bid/project", response_model=dict, dependencies=[Depends(verify_api_key)])
def create_bid_project(project: BidProjectCreate):
    """创建标书项目（需要 API Key）"""
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            cursor.execute(
                """
                INSERT INTO bid_projects
                    (title, source_file_name, file_path, parsed_data, status)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id, title, source_file_name, file_path, parsed_data,
                          status, created_at, updated_at
                """,
                (
                    project.title,
                    project.source_file_name,
                    project.file_path,
                    json.dumps(project.parsed_data) if project.parsed_data else None,
                    project.status,
                ),
            )
            row = cursor.fetchone()
            conn.commit()
            cursor.close()

        return {"status": "success", "data": dict(row)}

    except Exception as e:
        logger.error(f"创建标书项目失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/bid/projects", response_model=dict, dependencies=[Depends(verify_api_key)])
def list_bid_projects(
    status: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
):
    """
    标书项目列表
    修复 #7：添加分页支持
    """
    if not db_pool:
        return {"status": "error", "code": 503, "message": "数据库未连接", "detail": "DATABASE_URL 未设置"}
    try:
        # 分页参数校验
        if page < 1:
            page = 1
        if page_size < 1:
            page_size = 1
        if page_size > 100:
            page_size = 100
        offset = (page - 1) * page_size

        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            base_query = "FROM bid_projects"
            where_clause = ""
            params: list = []

            if status:
                where_clause = "WHERE status = %s"
                params.append(status)

            # 查询总数
            count_query = f"SELECT COUNT(*) as total {base_query} {where_clause}"
            cursor.execute(count_query, params)
            total = cursor.fetchone()["total"]

            # 查询分页数据
            data_query = f"""
                SELECT id, title, source_file_name, file_path, parsed_data,
                       status, created_at, updated_at
                {base_query} {where_clause}
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
            """
            cursor.execute(data_query, params + [page_size, offset])
            rows = cursor.fetchall()
            cursor.close()

        total_pages = (total + page_size - 1) // page_size if total > 0 else 0

        return {
            "status": "success",
            "count": len(rows),
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": total_pages,
            "data": [dict(r) for r in rows],
        }

    except Exception as e:
        logger.error(f"查询标书项目列表失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/bid/project/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def get_bid_project(project_id: int):
    """查询单个标书项目"""
    if not db_pool:
        return {"status": "error", "code": 503, "message": "数据库未连接", "detail": "DATABASE_URL 未设置"}
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

            cursor.execute(
                """
                SELECT id, title, source_file_name, file_path, parsed_data,
                       status, created_at, updated_at
                FROM bid_projects
                WHERE id = %s
                """,
                (project_id,),
            )
            row = cursor.fetchone()
            cursor.close()

        if not row:
            raise HTTPException(status_code=404, detail=f"项目 ID {project_id} 不存在")

        return {"status": "success", "data": dict(row)}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询标书项目失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================================
# 标书 AI 生成系统 — Day 2 新增 API
# =====================================================================

from fastapi import UploadFile, File
import shutil
import tempfile


# ==================== 任务 1: 文档解析 API ====================

@app.post("/api/bid/upload", response_model=dict, dependencies=[Depends(verify_api_key)])
async def upload_bid_document(file: UploadFile = File(...)):
    """
    上传招标文件（docx/pdf）并解析。

    - 接收文件并保存到 /tmp/
    - 调用 doc_parser.parse_file() 解析
    - 提取关键信息
    - 在 bid_projects 表中创建记录
    - 返回解析结果
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    # 验证文件类型
    allowed_ext = [".docx", ".pdf"]
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    if suffix not in allowed_ext:
        return error_response(400, "不支持的文件格式", f"仅支持 {', '.join(allowed_ext)}")

    # 保存到临时文件
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, prefix="bid_upload_", delete=False, dir="/tmp/") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        logger.info(f"文件已保存到: {tmp_path} ({len(content)} bytes)")

        # 解析文件
        from doc_parser import parse_and_extract
        result = parse_and_extract(tmp_path)

        # 在 bid_projects 表中创建记录
        parsed_data = result.get("parsed", {})
        key_info = result.get("key_info", {})

        # 合并 parsed 和 key_info 存储
        full_parsed = {**parsed_data, "key_info": key_info}

        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                """
                INSERT INTO bid_projects
                    (title, source_file_name, file_path, parsed_data, status)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id, title, source_file_name, file_path, parsed_data,
                          status, created_at, updated_at
                """,
                (
                    key_info.get("project_name") or filename,
                    filename,
                    tmp_path,
                    json.dumps(full_parsed, ensure_ascii=False),
                    "draft",
                ),
            )
            row = cursor.fetchone()
            conn.commit()
            cursor.close()

        return {
            "status": "success",
            "message": "文件解析成功",
            "file_path": tmp_path,
            "project_id": row["id"],
            "key_info": key_info,
            "parsed_summary": {
                "file_type": parsed_data.get("file_type"),
                "page_count": parsed_data.get("page_count"),
                "structure_items": len(parsed_data.get("structure", [])),
                "tables_count": len(parsed_data.get("tables", [])),
            },
            "data": dict(row),
        }

    except Exception as e:
        logger.error(f"文件上传解析失败: {e}")
        return error_response(500, "文件解析失败", str(e))


@app.get("/api/bid/analysis/{project_id}", dependencies=[Depends(verify_api_key)])
def get_bid_analysis(project_id: int):
    """
    获取招标文件解析结果。
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                """
                SELECT id, title, source_file_name, file_path, parsed_data,
                       status, created_at, updated_at
                FROM bid_projects
                WHERE id = %s
                """,
                (project_id,),
            )
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        parsed_data = row["parsed_data"]
        if isinstance(parsed_data, str):
            parsed_data = json.loads(parsed_data)

        key_info = parsed_data.get("key_info", {})

        return {
            "status": "success",
            "data": {
                "project_id": row["id"],
                "title": row["title"],
                "source_file_name": row["source_file_name"],
                "file_path": row["file_path"],
                "status": row["status"],
                "created_at": str(row["created_at"]),
                "key_info": key_info,
                "full_parsed": parsed_data,
            },
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取解析结果失败: {e}")
        return error_response(500, "获取解析结果失败", str(e))


# ==================== 任务 2: 动态模板引擎 API ====================

@app.post("/api/bid/generate-templates/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def generate_bid_templates(project_id: int):
    """
    为指定项目生成三套模板（报价/商务/技术）。
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 获取项目信息
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                "SELECT id, title, parsed_data FROM bid_projects WHERE id = %s",
                (project_id,),
            )
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        parsed_data = row["parsed_data"]
        if isinstance(parsed_data, str):
            parsed_data = json.loads(parsed_data)
        elif parsed_data is None:
            parsed_data = {}

        project_info = {
            "id": row["id"],
            "title": row["title"],
            "project_name": parsed_data.get("key_info", {}).get("project_name", row["title"]),
            "bid_number": parsed_data.get("key_info", {}).get("bid_number"),
        }

        # 推断项目类型
        from bid_template_engine import infer_project_type, generate_template_content, save_templates_to_db

        project_type = infer_project_type(parsed_data)
        logger.info(f"推断项目类型: {project_type}")

        # 生成三套模板
        templates = []
        for t_type in ["报价", "商务", "技术"]:
            tpl_content = generate_template_content(t_type, project_info)
            templates.append(tpl_content)

        # 保存到数据库
        saved = save_templates_to_db(project_id, templates, get_db_connection)

        return {
            "status": "success",
            "message": f"已生成 {len(saved)} 套模板",
            "project_type": project_type,
            "templates": saved,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成模板失败: {e}")
        return error_response(500, "生成模板失败", str(e))


@app.get("/api/bid/templates/{project_id}", dependencies=[Depends(verify_api_key)])
def get_bid_templates(project_id: int):
    """
    获取项目已生成的模板列表。
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        from bid_template_engine import load_templates_from_db

        templates = load_templates_from_db(project_id, get_db_connection)

        # 简化返回（去掉完整的 placeholders 列表）
        simplified = []
        for tpl in templates:
            content = tpl.get("content", {})
            simplified.append({
                "id": tpl["id"],
                "template_type": tpl["template_type"],
                "template_name": tpl["template_name"],
                "file_path": content.get("file_path"),
                "placeholder_count": len(content.get("placeholders", [])),
                "generated_at": content.get("generated_at"),
                "created_at": tpl["created_at"],
            })

        return {
            "status": "success",
            "count": len(simplified),
            "data": simplified,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取模板失败: {e}")
        return error_response(500, "获取模板失败", str(e))


# ==================== 任务 3: 资料需求分析 API ====================

@app.get("/api/bid/requirements/{project_id}", dependencies=[Depends(verify_api_key)])
def get_bid_requirements(project_id: int):
    """
    获取项目的资料需求清单。

    - 扫描模板中的占位符
    - 按类别分组
    - 标记已填充/待补充
    - 生成结构化需求清单
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 获取项目解析数据
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                "SELECT id, parsed_data FROM bid_projects WHERE id = %s",
                (project_id,),
            )
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        parsed_data = row["parsed_data"]
        if isinstance(parsed_data, str):
            parsed_data = json.loads(parsed_data)
        elif parsed_data is None:
            parsed_data = {}

        from requirements_analyzer import analyze_requirements

        result = analyze_requirements(project_id, parsed_data, get_db_connection)

        return {
            "status": "success",
            "data": result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取需求清单失败: {e}")
        return error_response(500, "获取需求清单失败", str(e))


# ==================== 任务 4: 数据匹配/自动填充 API ====================

import uuid
from pydantic import BaseModel

class AutoFillRequest(BaseModel):
    bid_amount: Optional[float] = None


class BidMergeRequest(BaseModel):
    company_name: Optional[str] = None
    contact_person: Optional[str] = None
    contact_phone: Optional[str] = None
    bid_amount: Optional[float] = 100000.0
    project_type: Optional[str] = "物联网卡"
    custom_fields: Optional[Dict[str, Any]] = None


@app.post("/api/bid/auto-fill/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def auto_fill_bid(project_id: int, request: AutoFillRequest = AutoFillRequest()):
    """
    自动填充标书模板。

    - 查询公司资料
    - 查询历史业绩
    - 自动匹配占位符
    - 生成填充后的模板文件
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        from data_matcher import auto_fill_project

        result = auto_fill_project(
            project_id=project_id,
            get_db_connection_func=get_db_connection,
            bid_amount=request.bid_amount,
        )

        if "error" in result:
            return error_response(404, result["error"], "")

        return {
            "status": "success",
            "data": result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"自动填充失败: {e}")
        return error_response(500, "自动填充失败", str(e))


# =====================================================================
# 标书校验 & 评分 & 飞书通知 — Day 4 新增 API
# =====================================================================

from disqualification_checker import check_disqualification
from scoring_report import generate_scoring_report
from feishu_notifier import notify_completion


class NotifyRequest(BaseModel):
    webhook_url: Optional[str] = None


@app.get("/api/bid/disqualification/{project_id}", response_model=dict,
         dependencies=[Depends(verify_api_key)])
def get_disqualification_result(project_id: int):
    """
    获取废标项检查结果。

    检查项：
    - 关键资质缺失（营业执照、资质证书等）
    - 报价超过预算上限
    - 必填项未完成
    - 投标有效期不足
    - 法定代表人或授权代表缺失
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 确认项目存在
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM bid_projects WHERE id = %s", (project_id,))
            exists = cursor.fetchone()
            cursor.close()

        if not exists:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        result = check_disqualification(project_id, get_db_connection=get_db_connection)

        return {
            "status": "success",
            "data": result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"废标检查失败: {e}")
        return error_response(500, "废标检查失败", str(e))


@app.get("/api/bid/scoring-report/{project_id}", response_model=dict,
         dependencies=[Depends(verify_api_key)])
def get_scoring_report(project_id: int):
    """
    获取评分覆盖率报告。

    评分维度：
    - 技术方案 (30%)
    - 项目管理 (15%)
    - 业绩经验 (15%)
    - 团队资质 (10%)
    - 质量保证 (10%)
    - 价格 (20%)
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 确认项目存在
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM bid_projects WHERE id = %s", (project_id,))
            exists = cursor.fetchone()
            cursor.close()

        if not exists:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        result = generate_scoring_report(project_id, get_db_connection=get_db_connection)

        return {
            "status": "success",
            "data": result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"评分报告生成失败: {e}")
        return error_response(500, "评分报告生成失败", str(e))


@app.post("/api/bid/notify/{project_id}", response_model=dict,
          dependencies=[Depends(verify_api_key)])
def trigger_feishu_notify(project_id: int, request: NotifyRequest = NotifyRequest()):
    """
    手动触发飞书通知。

    通知内容：
    - 项目名称
    - 完成时间
    - 废标项检查结果
    - 预估得分
    - 下载链接
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 确认项目存在
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, title FROM bid_projects WHERE id = %s", (project_id,))
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        # 构建 base URL（从环境变量或请求头推断）
        base_url = os.environ.get("APP_BASE_URL", "")

        success = notify_completion(
            project_id=project_id,
            webhook_url=request.webhook_url,
            get_db_connection=get_db_connection,
            base_url=base_url,
        )

        if success:
            return {
                "status": "success",
                "message": f"已向飞书发送项目 '{row[1]}' 的完成通知",
            }
        else:
            return error_response(500, "飞书通知发送失败", "请检查 Webhook URL 配置")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"飞书通知触发失败: {e}")
        return error_response(500, "飞书通知触发失败", str(e))


# =====================================================================
# 标书 AI 内容生成 — Day 3 新增 API（P0 修复：接入路由）
# =====================================================================

from ai_generator import (
    generate_technical_solution,
    generate_project_understanding,
    generate_work_plan,
    generate_performance_guarantee,
    generate_service_commitment,
    generate_all,
    record_generation_start,
    record_generation_complete,
    get_generation_status,
    list_generations,
)


class AIGenerateRequest(BaseModel):
    """AI 生成请求"""
    modules: Optional[List[str]] = None  # 指定模块，不传则全量生成


@app.post("/api/bid/ai-generate/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def ai_generate_endpoint(project_id: int, request: AIGenerateRequest = AIGenerateRequest()):
    """
    触发 AI 生成标书内容。

    支持 5 个模块：
    - technical_solution: 技术方案
    - project_understanding: 项目理解
    - work_plan: 工作规划
    - performance_guarantee: 履约保障
    - service_commitment: 服务承诺

    不传 modules 参数则触发全量生成（5 个模块）。
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 查询项目信息
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                "SELECT id, title, parsed_data FROM bid_projects WHERE id = %s",
                (project_id,),
            )
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        parsed_data = row["parsed_data"]
        if isinstance(parsed_data, str):
            parsed_data = json.loads(parsed_data)
        elif parsed_data is None:
            return error_response(400, "项目解析数据为空", "请先上传并解析招标文件")

        # 记录生成开始
        all_modules = [
            "technical_solution",
            "project_understanding",
            "work_plan",
            "performance_guarantee",
            "service_commitment",
        ]
        modules = request.modules if request.modules else all_modules
        task_id = record_generation_start(project_id, modules)

        logger.info(f"AI 生成任务开始: task_id={task_id}, project_id={project_id}, modules={modules}")

        # 执行生成
        result = generate_all(project_id, parsed_data, get_db_connection)

        # 记录生成完成
        record_generation_complete(task_id, result)

        return {
            "status": "success",
            "message": "AI 生成完成",
            "task_id": task_id,
            "data": result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI 生成失败: {e}")
        return error_response(500, "AI 生成失败", str(e))


@app.get("/api/bid/ai-status/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def ai_status_endpoint(project_id: int):
    """
    查询项目 AI 生成状态和结果。

    返回该项目的最近一次生成任务状态。
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        # 查询项目是否存在
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                "SELECT id, title, status FROM bid_projects WHERE id = %s",
                (project_id,),
            )
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        # 查询生成状态
        generations = list_generations(project_id=project_id)

        if not generations:
            return {
                "status": "success",
                "data": {
                    "project_id": project_id,
                    "project_title": row["title"],
                    "project_status": row["status"],
                    "generation_status": "no_task",
                    "message": "暂无 AI 生成任务",
                },
            }

        # 返回最近一次任务
        latest = generations[0]
        return {
            "status": "success",
            "data": {
                "project_id": project_id,
                "project_title": row["title"],
                "project_status": row["status"],
                "task_id": latest.get("task_id"),
                "generation_status": latest.get("status"),
                "started_at": latest.get("started_at"),
                "completed_at": latest.get("completed_at"),
                "elapsed_seconds": latest.get("elapsed_seconds"),
                "modules": latest.get("modules"),
                "results": latest.get("results", {}),
            },
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询 AI 生成状态失败: {e}")
        return error_response(500, "查询 AI 生成状态失败", str(e))


# =====================================================================
# 标书整合引擎 — Day 4 新增 API
# =====================================================================

@app.post("/api/bid/merge/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def merge_bid(project_id: int, request: BidMergeRequest = BidMergeRequest()):
    """
    合并三套标书为完整文档。

    - 添加统一封面页
    - 自动生成目录
    - 统一页眉页脚
    - 页码连续编号
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        from bid_merger import merge_bid_documents_to_default

        # 获取项目信息
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute("SELECT id, title FROM bid_projects WHERE id = %s", (project_id,))
            row = cursor.fetchone()
            cursor.close()

        if not row:
            return error_response(404, "项目不存在", f"项目 ID {project_id} 不存在")

        # 获取公司资料
        company_name = request.company_name
        contact_person = request.contact_person
        contact_phone = request.contact_phone

        if not company_name or not contact_person or not contact_phone:
            try:
                with get_db_connection() as conn:
                    cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
                    cursor.execute(
                        "SELECT company_name, contact_person, phone FROM company_profiles ORDER BY created_at DESC LIMIT 1"
                    )
                    comp_row = cursor.fetchone()
                    cursor.close()
                if comp_row:
                    company_name = company_name or comp_row["company_name"]
                    contact_person = contact_person or comp_row["contact_person"]
                    contact_phone = contact_phone or comp_row["phone"]
            except Exception as e:
                logger.warning(f"获取公司资料失败: {e}")

        company_name = company_name or "默认投标公司"
        contact_person = contact_person or "张三"
        contact_phone = contact_phone or "13800138000"

        # 执行合并
        result = merge_bid_documents_to_default(
            project_id=project_id,
            company_name=company_name,
            contact_person=contact_person,
            contact_phone=contact_phone,
            bid_amount=request.bid_amount,
            project_type=request.project_type,
            custom_fields=request.custom_fields,
            get_db_connection_func=get_db_connection,
        )

        # 更新项目状态
        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "UPDATE bid_projects SET status = 'merged', updated_at = CURRENT_TIMESTAMP WHERE id = %s",
                    (project_id,)
                )
                conn.commit()
                cursor.close()
        except Exception as e:
            logger.warning(f"更新项目状态失败: {e}")

        return {
            "status": "success",
            "message": "标书合并完成",
            "data": result,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"合并标书失败: {e}")
        return error_response(500, "合并标书失败", str(e))


@app.get("/api/bid/merged/{project_id}", dependencies=[Depends(verify_api_key)])
def get_merged_bid(project_id: int):
    """
    获取已合并的标书信息。
    查找最新的合并文件。
    """
    try:
        from pathlib import Path
        output_dir = Path("/tmp/bid-outputs") / str(project_id)

        if not output_dir.exists():
            return error_response(404, "未找到合并文件", f"项目 {project_id} 没有合并输出")

        # 查找最新的 docx 文件
        docx_files = list(output_dir.glob("*.docx"))
        if not docx_files:
            return error_response(404, "未找到合并文件", f"项目 {project_id} 没有 .docx 输出")

        latest = max(docx_files, key=lambda f: f.stat().st_mtime)

        return {
            "status": "success",
            "data": {
                "project_id": project_id,
                "file_path": str(latest),
                "file_name": latest.name,
                "file_size_bytes": latest.stat().st_size,
                "modified_at": datetime.fromtimestamp(latest.stat().st_mtime).isoformat(),
            },
        }

    except Exception as e:
        logger.error(f"获取合并标书失败: {e}")
        return error_response(500, "获取合并标书失败", str(e))


# =====================================================================
# PDF 导出 — Day 4 新增 API
# =====================================================================

@app.post("/api/bid/export-pdf/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def export_bid_pdf(project_id: int):
    """
    将合并后的标书导出为 PDF。

    转换策略:
      1. LibreOffice (headless) — 推荐
      2. pandoc — 备选
      3. weasyprint — 备选
      4. docx2pdf — 备选
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        from pathlib import Path
        from pdf_exporter import export_project_pdf, get_available_converters

        output_dir = Path("/tmp/bid-outputs") / str(project_id)

        if not output_dir.exists():
            return error_response(404, "未找到标书文件",
                                   f"请先调用 /api/bid/merge/{project_id} 合并标书")

        # 查找最新的 docx 文件
        docx_files = list(output_dir.glob("*.docx"))
        if not docx_files:
            return error_response(404, "未找到 docx 文件",
                                   f"项目 {project_id} 没有 .docx 输出")

        latest_docx = max(docx_files, key=lambda f: f.stat().st_mtime)

        # 检查转换器
        converters = get_available_converters()
        available = [c for c in converters if c["status"] == "available"]
        if not available:
            return {
                "status": "warning",
                "message": "当前服务器没有可用的 PDF 转换工具",
                "converters": converters,
                "hint": "请安装 LibreOffice: apt install libreoffice",
            }

        # 执行转换
        result = export_project_pdf(project_id, str(latest_docx))

        return {
            "status": "success",
            "message": "PDF 导出完成",
            "data": result,
            "converters": converters,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"导出 PDF 失败: {e}")
        return error_response(500, "导出 PDF 失败", str(e))


@app.get("/api/bid/pdf/{project_id}", dependencies=[Depends(verify_api_key)])
def get_bid_pdf(project_id: int):
    """
    获取项目的 PDF 文件信息。
    """
    try:
        from pathlib import Path
        output_dir = Path("/tmp/bid-outputs") / str(project_id)

        if not output_dir.exists():
            return error_response(404, "未找到文件", f"项目 {project_id} 没有输出文件")

        pdf_files = list(output_dir.glob("*.pdf"))
        if not pdf_files:
            return error_response(404, "未找到 PDF 文件",
                                   f"请先调用 /api/bid/export-pdf/{project_id} 导出 PDF")

        latest = max(pdf_files, key=lambda f: f.stat().st_mtime)

        return {
            "status": "success",
            "data": {
                "project_id": project_id,
                "file_path": str(latest),
                "file_name": latest.name,
                "file_size_bytes": latest.stat().st_size,
                "modified_at": datetime.fromtimestamp(latest.stat().st_mtime).isoformat(),
            },
        }

    except Exception as e:
        logger.error(f"获取 PDF 失败: {e}")
        return error_response(500, "获取 PDF 失败", str(e))


@app.get("/api/bid/converters", dependencies=[Depends(verify_api_key)])
def list_converters():
    """
    列出当前可用的 PDF 转换工具。
    """
    try:
        from pdf_exporter import get_available_converters
        converters = get_available_converters()
        return {
            "status": "success",
            "data": converters,
        }
    except Exception as e:
        logger.error(f"获取转换器列表失败: {e}")
        return error_response(500, "获取转换器列表失败", str(e))


# =====================================================================
# 下载链接生成 — Day 4 新增 API
# =====================================================================

@app.post("/api/bid/download/{project_id}", response_model=dict, dependencies=[Depends(verify_api_key)])
def generate_download_link(project_id: int, file_type: str = "docx"):
    """
    生成临时下载链接（24 小时有效）。

    - file_type: docx 或 pdf
    """
    if not db_pool:
        return error_response(503, "数据库未连接", "DATABASE_URL 未设置")

    try:
        from datetime import timedelta
        from pathlib import Path

        output_dir = Path("/tmp/bid-outputs") / str(project_id)

        if not output_dir.exists():
            return error_response(404, "未找到文件", f"项目 {project_id} 没有输出文件")

        # 查找文件
        ext = ".pdf" if file_type == "pdf" else ".docx"
        files = list(output_dir.glob(f"*{ext}"))

        if not files:
            return error_response(404, f"未找到 {ext} 文件",
                                   f"项目 {project_id} 没有 {ext} 输出。" +
                                   (f"请先调用 /api/bid/export-pdf/{project_id} 导出 PDF" if file_type == "pdf" else "请先调用 /api/bid/merge/{project_id} 合并标书"))

        latest = max(files, key=lambda f: f.stat().st_mtime)

        # 生成 token
        download_token = str(uuid.uuid4())
        expires_at = datetime.now() + timedelta(hours=24)

        # 存储到数据库
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                INSERT INTO bid_downloads
                    (download_token, project_id, file_path, file_type, expires_at)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id
                """,
                (download_token, project_id, str(latest), file_type, expires_at)
            )
            conn.commit()
            cursor.close()

        # 构建下载 URL
        base_url = os.environ.get("BASE_URL", "")
        download_url = f"{base_url}/api/bid/download/{download_token}"

        return {
            "status": "success",
            "message": "下载链接已生成（24 小时有效）",
            "data": {
                "download_url": download_url,
                "token": download_token,
                "file_name": latest.name,
                "file_size_bytes": latest.stat().st_size,
                "file_type": file_type,
                "expires_at": expires_at.isoformat(),
            },
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成下载链接失败: {e}")
        return error_response(500, "生成下载链接失败", str(e))


@app.get("/api/bid/download/{token}")
def download_file(token: str):
    """
    通过 token 下载文件。
    - 验证 token 是否有效（未过期）
    - 增加下载次数
    - 返回文件流
    """
    if not db_pool:
        raise HTTPException(status_code=503, detail="数据库未连接")

    try:
        from fastapi.responses import FileResponse

        # 查询 token
        with get_db_connection() as conn:
            cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            cursor.execute(
                """
                SELECT id, download_token, project_id, file_path, file_type,
                       expires_at, download_count
                FROM bid_downloads
                WHERE download_token = %s
                """,
                (token,)
            )
            row = cursor.fetchone()

            if not row:
                cursor.close()
                raise HTTPException(status_code=404, detail="下载链接不存在")

            if row["expires_at"] < datetime.now():
                cursor.close()
                raise HTTPException(status_code=410, detail="下载链接已过期")

            # 更新下载次数
            cursor.execute(
                """
                UPDATE bid_downloads
                SET download_count = download_count + 1,
                    last_downloaded_at = CURRENT_TIMESTAMP
                WHERE id = %s
                """,
                (row["id"],)
            )
            conn.commit()
            cursor.close()

        file_path = row["file_path"]
        if not Path(file_path).exists():
            raise HTTPException(status_code=404, detail=f"文件不存在: {file_path}")

        file_name = Path(file_path).name
        media_type = (
            "application/pdf" if row["file_type"] == "pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        return FileResponse(
            path=file_path,
            media_type=media_type,
            filename=file_name,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载文件失败: {e}")
        raise HTTPException(status_code=500, detail=f"下载失败: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
