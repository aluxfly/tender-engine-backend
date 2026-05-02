"""
标书 AI 生成系统 — Day 2: 数据匹配引擎
=======================================
功能:
  1. 查询 company_profiles 表的公司资料
  2. 查询历史业绩数据
  3. 查询团队人员信息
  4. 自动匹配并填充模板中的占位符
  5. 更新 bid_fill_status 表
"""

import json
import re
import logging
import os
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent / "templates"

# ==================== 数据查询 ====================

def get_company_profile(get_db_connection_func) -> Optional[Dict[str, Any]]:
    """获取公司资料。"""
    with get_db_connection_func() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, company_name, credit_code, legal_representative,
                   contact_person, phone, email, address, bank_info, qualifications
            FROM company_profiles
            ORDER BY updated_at DESC
            LIMIT 1
            """
        )
        row = cursor.fetchone()
        cursor.close()

    if not row:
        return None

    return {
        "id": row[0],
        "company_name": row[1],
        "credit_code": row[2],
        "legal_representative": row[3],
        "contact_person": row[4],
        "phone": row[5],
        "email": row[6],
        "address": row[7],
        "bank_info": json.loads(row[8]) if row[8] and isinstance(row[8], str) else (row[8] or {}),
        "qualifications": json.loads(row[9]) if row[9] and isinstance(row[9], str) else (row[9] or []),
    }


def get_project_performance(project_id: int, get_db_connection_func) -> List[Dict[str, Any]]:
    """获取历史业绩数据。"""
    with get_db_connection_func() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, title, parsed_data, status, created_at
            FROM bid_projects
            WHERE id != %s AND status = 'completed'
            ORDER BY created_at DESC
            LIMIT 10
            """,
            (project_id,),
        )
        rows = cursor.fetchall()
        cursor.close()

    results = []
    for row in rows:
        parsed = row[2]
        if parsed and isinstance(parsed, str):
            parsed = json.loads(parsed)
        results.append({
            "id": row[0],
            "title": row[1],
            "parsed_data": parsed or {},
            "status": row[3],
            "created_at": str(row[4]),
        })

    return results


def get_project_info(project_id: int, get_db_connection_func) -> Optional[Dict[str, Any]]:
    """获取项目信息。"""
    with get_db_connection_func() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, title, source_file_name, file_path, parsed_data, status
            FROM bid_projects
            WHERE id = %s
            """,
            (project_id,),
        )
        row = cursor.fetchone()
        cursor.close()

    if not row:
        return None

    parsed = row[4]
    if parsed and isinstance(parsed, str):
        parsed = json.loads(parsed)

    return {
        "id": row[0],
        "title": row[1],
        "source_file_name": row[2],
        "file_path": row[3],
        "parsed_data": parsed or {},
        "status": row[5],
    }


# ==================== 占位符映射 ====================

# 占位符名称 → 数据源的映射规则
PLACEHOLDER_MAP = {
    # 公司信息
    "投标人名称": ("company_profile", "company_name"),
    "投标人名称": ("company_profile", "company_name"),
    "公司名称": ("company_profile", "company_name"),
    "企业名称": ("company_profile", "company_name"),
    "统一社会信用代码": ("company_profile", "credit_code"),
    "信用代码": ("company_profile", "credit_code"),
    "代码": ("company_profile", "credit_code"),
    "法定代表人": ("company_profile", "legal_representative"),
    "法定代表人姓名": ("company_profile", "legal_representative"),
    "联系人": ("company_profile", "contact_person"),
    "联系电话": ("company_profile", "phone"),
    "电话号码": ("company_profile", "phone"),
    "地址": ("company_profile", "address"),
    "注册地址": ("company_profile", "address"),
    "投标人地址": ("company_profile", "address"),
    "邮箱": ("company_profile", "email"),
    "邮箱地址": ("company_profile", "email"),
    "电子邮箱": ("company_profile", "email"),
    # 银行信息
    "银行名称": ("company_profile", "bank_info.bank_name"),
    "开户银行": ("company_profile", "bank_info.bank_name"),
    "开户账号": ("company_profile", "bank_info.account"),
    "银行账号": ("company_profile", "bank_info.account"),
    "账号": ("company_profile", "bank_info.account"),
    "资信等级": ("company_profile", "bank_info.credit_rating"),
    # 项目信息
    "项目名称": ("project", "parsed_data.project_name"),
    "招标编号": ("project", "parsed_data.bid_number"),
    "采购人": ("project", "parsed_data.buyer"),
    "招标人名称": ("project", "parsed_data.buyer"),
    "截止日期": ("project", "parsed_data.deadline"),
    "预算": ("project", "parsed_data.budget"),
    # 报价
    "人民币大写金额": ("project", "bid_amount_text"),
    "小写金额": ("project", "bid_amount"),
    "投标总报价": ("project", "bid_amount"),
    # 授权
    "授权代表姓名": ("company_profile", "contact_person"),
    "被授权人": ("company_profile", "contact_person"),
}


def resolve_value(source_type: str, source_path: str, company: Dict,
                  project: Dict, performances: List) -> Optional[str]:
    """解析数据源路径，获取实际值。"""
    if source_type == "company_profile":
        parts = source_path.split(".")
        value = company
        for part in parts:
            if isinstance(value, dict):
                value = value.get(part)
            else:
                return None
        return str(value) if value else None

    elif source_type == "project":
        if source_path.startswith("parsed_data."):
            key = source_path[len("parsed_data."):]
            parsed = project.get("parsed_data", {})
            value = parsed.get(key)
            return str(value) if value else None
        elif source_path == "bid_amount_text":
            # 数字转中文大写
            amount = project.get("bid_amount")
            if amount:
                return number_to_chinese_uppercase(amount)
            return None
        elif source_path == "bid_amount":
            amount = project.get("bid_amount")
            return str(amount) if amount else None
        return None

    return None


def number_to_chinese_uppercase(number) -> str:
    """数字转中文大写金额。"""
    if not number:
        return "[人民币大写金额]"

    try:
        num = float(number)
    except (ValueError, TypeError):
        return str(number)

    if num < 0:
        return f"负{number_to_chinese_uppercase(abs(num))}"

    integer_part = int(num)
    decimal_part = round((num - integer_part) * 100)

    digits = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿"]

    if integer_part == 0:
        int_part_str = "零"
    else:
        # 按4位一组拆分
        num_str = str(integer_part)
        total_len = len(num_str)
        int_part_str = ""
        prev_zero = False
        # 记录哪些数量级已经添加过万/亿
        added_big_units = set()

        for i, d in enumerate(num_str):
            dv = int(d)
            pos_from_right = total_len - 1 - i
            big_unit_idx = pos_from_right // 4  # 0=个级, 1=万级, 2=亿级

            if dv == 0:
                if pos_from_right > 0:
                    prev_zero = True
            else:
                if prev_zero:
                    int_part_str += "零"
                prev_zero = False
                # 先输出数字+单位
                int_part_str += digits[dv] + units[pos_from_right % 4]
                # 检查当前位置之后（更低的位置）是否属于更低数量级
                # 如果是该数量级的第一个非零数字且该级单位未加过，则加万/亿
                # 简化判断：只有当 pos_from_right % 4 == 0（即万/亿位）时立即加
                # 或者当后续所有数字都属于更低数量级时
                remaining_len = pos_from_right  # 当前位之后还有多少位
                if big_unit_idx > 0:
                    # 检查是否应该在此处加大单位
                    # 如果当前位的 pos_from_right % 4 == 0（正好是万/亿的起始位），直接加
                    # 否则，检查是否该数量级还没有被标记
                    if pos_from_right % 4 == 0 and big_unit_idx not in added_big_units:
                        int_part_str += big_units[big_unit_idx]
                        added_big_units.add(big_unit_idx)
                    elif remaining_len > 0 and remaining_len <= big_unit_idx * 4:
                        # 当前位之后都属于更低数量级，且该级未加过大单位
                        if big_unit_idx not in added_big_units:
                            int_part_str += big_units[big_unit_idx]
                            added_big_units.add(big_unit_idx)

    # 小数部分
    if decimal_part == 0:
        return f"{int_part_str}元整"
    else:
        jiao = decimal_part // 10
        fen = decimal_part % 10
        result = f"{int_part_str}元"
        if jiao > 0:
            result += f"{digits[jiao]}角"
        if fen > 0:
            result += f"{digits[fen]}分"
        return result


# ==================== 模板填充 ====================

def find_template_files(project_id: int) -> List[Tuple[str, str]]:
    """查找项目相关的模板文件。"""
    templates = []
    for f in TEMPLATES_DIR.iterdir():
        if f.suffix == ".docx" and f"template_" in f.name:
            # 检查是否属于该项目
            if f.name.startswith(f"template_") and str(project_id) in f.name:
                for t_type in ["报价", "商务", "技术"]:
                    if t_type in f.name:
                        templates.append((t_type, str(f)))

    # 如果没有找到特定项目的模板，使用默认模板
    if not templates:
        default_map = {
            "报价": "bid-price-file.docx",
            "商务": "bid-business-file.docx",
            "技术": "bid-technical-file.docx",
        }
        for t_type, fname in default_map.items():
            fpath = TEMPLATES_DIR / fname
            if fpath.exists():
                templates.append((t_type, str(fpath)))

    return templates


def fill_template(input_path: str, output_path: str, replacements: Dict[str, str]) -> Dict[str, Any]:
    """
    填充模板文件中的占位符。

    返回:
    {
        "input_path": "...",
        "output_path": "...",
        "replacements_made": [...],
        "unfilled_placeholders": [...],
    }
    """
    doc = Document(input_path)
    replacements_made = []
    unfilled = []

    # 替换段落中的占位符
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if new and f"[{old}]" in para.text:
                for run in para.runs:
                    if f"[{old}]" in run.text:
                        run.text = run.text.replace(f"[{old}]", str(new))
                        replacements_made.append({"placeholder": old, "value": str(new)})

        # 收集未填充的占位符
        for match in re.finditer(r"\[([^\]]+)\]", para.text):
            ph_name = match.group(1).strip()
            if ph_name not in replacements or not replacements.get(ph_name):
                unfilled.append(ph_name)

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if new and f"[{old}]" in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if f"[{old}]" in run.text:
                                    run.text = run.text.replace(f"[{old}]", str(new))
                                    replacements_made.append({"placeholder": old, "value": str(new)})

                for match in re.finditer(r"\[([^\]]+)\]", cell.text):
                    ph_name = match.group(1).strip()
                    if ph_name not in replacements or not replacements.get(ph_name):
                        unfilled.append(ph_name)

    doc.save(output_path)

    # 去重
    seen_unfilled = set()
    unique_unfilled = []
    for u in unfilled:
        if u not in seen_unfilled:
            seen_unfilled.add(u)
            unique_unfilled.append(u)

    return {
        "input_path": input_path,
        "output_path": output_path,
        "replacements_made": replacements_made,
        "unfilled_placeholders": unique_unfilled,
        "replacement_count": len(replacements_made),
        "unfilled_count": len(unique_unfilled),
    }


# ==================== 主流程 ====================

def auto_fill_project(project_id: int, get_db_connection_func,
                      bid_amount: Optional[float] = None) -> Dict[str, Any]:
    """
    自动填充标书。

    返回:
    {
        "project_id": ...,
        "fill_time": "...",
        "templates_filled": [...],
        "summary": {...},
        "unfilled_items": [...],
    }
    """
    # 1. 获取项目信息
    project = get_project_info(project_id, get_db_connection_func)
    if not project:
        return {"error": f"项目 ID {project_id} 不存在"}

    # 2. 获取公司资料
    company = get_company_profile(get_db_connection_func)
    if not company:
        logger.warning("未找到公司资料，填充效果受限")
        company = {}

    # 3. 获取历史业绩
    performances = get_project_performance(project_id, get_db_connection_func)

    # 4. 构建替换映射
    replacements = {}
    fill_results = []

    # 从 PLACEHOLDER_MAP 填充
    for placeholder, (source_type, source_path) in PLACEHOLDER_MAP.items():
        value = resolve_value(source_type, source_path, company, project, performances)
        if value:
            replacements[placeholder] = value

    # 从公司资质填充
    if isinstance(company.get("qualifications"), list):
        quals = company["qualifications"]
        if quals:
            replacements["资质"] = ", ".join(str(q) for q in quals)
            replacements["资质等级"] = ", ".join(str(q) for q in quals)
            replacements["证书"] = ", ".join(str(q) for q in quals)

    # 从银行信息填充
    if isinstance(company.get("bank_info"), dict):
        for key, value in company["bank_info"].items():
            if value:
                replacements[key] = str(value)

    # 从项目解析数据填充
    parsed = project.get("parsed_data", {})
    if isinstance(parsed, dict):
        for key, value in parsed.items():
            if value and isinstance(value, str) and len(value) < 200:
                # 尝试匹配常见的占位符名称
                for ph_name in ["项目名称", "招标编号", "预算金额", "采购人", "招标代理",
                                "资质要求", "联系方式", "招标方式"]:
                    if key in ph_name.lower() or ph_name.lower() in key.lower():
                        if ph_name not in replacements:
                            replacements[ph_name] = str(value)

    # 5. 填充模板文件
    templates = find_template_files(project_id)
    all_unfilled = []

    output_dir = TEMPLATES_DIR / f"filled_{project_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    output_dir.mkdir(exist_ok=True)

    for t_type, input_path in templates:
        output_filename = f"filled_{t_type}_{project_id}.docx"
        output_path = str(output_dir / output_filename)

        fill_result = fill_template(input_path, output_path, replacements)
        fill_result["template_type"] = t_type
        fill_results.append(fill_result)
        all_unfilled.extend(fill_result.get("unfilled_placeholders", []))

    # 6. 更新 bid_fill_status 表
    try:
        with get_db_connection_func() as conn:
            cursor = conn.cursor()

            # 更新已填充的状态
            for ph, value in replacements.items():
                cursor.execute(
                    """
                    UPDATE bid_fill_status
                    SET fill_status = %s, fill_source = %s, value = %s, updated_at = CURRENT_TIMESTAMP
                    WHERE project_id = %s AND field_name = %s
                    """,
                    ("filled", "auto_fill", str(value), project_id, ph),
                )

            # 为未填充的创建记录
            for ph in set(all_unfilled):
                cursor.execute(
                    """
                    SELECT id FROM bid_fill_status
                    WHERE project_id = %s AND field_name = %s
                    """,
                    (project_id, ph),
                )
                if not cursor.fetchone():
                    cursor.execute(
                        """
                        INSERT INTO bid_fill_status (project_id, field_name, fill_status)
                        VALUES (%s, %s, 'unfilled')
                        """,
                        (project_id, ph),
                    )

            conn.commit()
            cursor.close()
    except Exception as e:
        logger.warning(f"更新填充状态失败: {e}")

    # 7. 构建返回结果
    total_replacements = sum(r["replacement_count"] for r in fill_results)
    total_unfilled = len(set(all_unfilled))

    return {
        "project_id": project_id,
        "fill_time": datetime.now().isoformat(),
        "templates_filled": fill_results,
        "output_directory": str(output_dir),
        "summary": {
            "total_replacements": total_replacements,
            "unique_placeholders_filled": len(replacements),
            "total_unfilled": total_unfilled,
            "fill_rate": round(
                len(replacements) / (len(replacements) + total_unfilled) * 100, 1
            ) if (len(replacements) + total_unfilled) > 0 else 0,
        },
        "unfilled_items": list(set(all_unfilled)),
        "company_profile_used": bool(company),
        "performances_used": len(performances),
    }
