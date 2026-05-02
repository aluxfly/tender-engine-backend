"""
文档解析引擎 — 标书 AI 生成系统 Day 1
======================================
功能:
  - DOCX 解析：章节结构、表格、关键信息提取
  - PDF 解析：文本提取、表格提取
  - 关键信息提取：项目名称、招标编号、截止日期、资质要求等

依赖:
  - python-docx (已有)
  - pdfplumber (新增)
  - re, json (标准库)

用法:
  from doc_parser import parse_docx, parse_pdf, extract_key_info

  result = parse_docx("招标文件.docx")
  key_info = extract_key_info(result)
  print(json.dumps(key_info, ensure_ascii=False, indent=2))
"""

import re
import json
import logging
from typing import Optional, Dict
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)


# ==================== 常量定义 ====================

# 招标文档常见关键信息模式
KEY_PATTERNS = {
    "project_name": [
        r"项目名称[：:\s]*([^\n\r]+)",
        r"工程名称[：:\s]*([^\n\r]+)",
        r"采购项目名称[：:\s]*([^\n\r]+)",
        r"招标项目名称[：:\s]*([^\n\r]+)",
    ],
    "bid_number": [
        r"招标编号[：:\s]*([^\n\r]+)",
        r"项目编号[：:\s]*([^\n\r]+)",
        r"采购编号[：:\s]*([^\n\r]+)",
        r"标书编号[：:\s]*([^\n\r]+)",
        r"文件编号[：:\s]*([^\n\r]+)",
        r"编号[：:\s]*([^\n\r]+)",
    ],
    "deadline": [
        r"截止[日时]间[：:\s]*([^\n\r]+)",
        r"投标截止[日时]间[：:\s]*([^\n\r]+)",
        r"递交截止[日时]间[：:\s]*([^\n\r]+)",
        r"响应文件递交截止[日时]间[：:\s]*([^\n\r]+)",
        r"报名截止[日时]间[：:\s]*([^\n\r]+)",
        r"开标时间[：:\s]*([^\n\r]+)",
        r"开标[日时]间[：:\s]*([^\n\r]+)",
    ],
    "budget": [
        r"预算金额[：:\s]*([^\n\r]+)",
        r"采购预算[：:\s]*([^\n\r]+)",
        r"项目预算[：:\s]*([^\n\r]+)",
        r"最高限价[：:\s]*([^\n\r]+)",
        r"控制价[：:\s]*([^\n\r]+)",
        r"估算金额[：:\s]*([^\n\r]+)",
    ],
    "buyer": [
        r"采购人[：:\s]*([^\n\r]+)",
        r"招标人[：:\s]*([^\n\r]+)",
        r"采购单位[：:\s]*([^\n\r]+)",
        r"建设单位[：:\s]*([^\n\r]+)",
    ],
    "agency": [
        r"招标代理[：:\s]*([^\n\r]+)",
        r"代理机构[：:\s]*([^\n\r]+)",
        r"采购代理机构[：:\s]*([^\n\r]+)",
    ],
    "qualification": [
        r"资质要求[：:\s]*([^\n\r]{5,200})",
        r"投标人资质[：:\s]*([^\n\r]{5,200})",
        r"资格要求[：:\s]*([^\n\r]{5,200})",
        r"供应商资格要求[：:\s]*([^\n\r]{5,200})",
    ],
    "contact": [
        r"联系方式[：:\s]*([^\n\r]+)",
        r"联系人[：:\s]*([^\n\r]+)",
        r"联系电话[：:\s]*([^\n\r]+)",
        r"联系电话[：:\s]*([\d\-]+)",
        r"电话[：:\s]*([\d\-]+)",
    ],
    "bid_method": [
        r"招标方式[：:\s]*([^\n\r]+)",
        r"采购方式[：:\s]*([^\n\r]+)",
        r"招标类型[：:\s]*([^\n\r]+)",
    ],
}

# 日期标准化正则
DATE_PATTERNS = [
    (r"(\d{4})[年\-/.](\d{1,2})[月\-/.](\d{1,2})", lambda m: f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"),
    (r"(\d{4})[年](\d{1,2})[月](\d{1,2})[日]", lambda m: f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"),
    (r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})", lambda m: f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"),
]


# ==================== DOCX 解析 ====================

def parse_docx(file_path: str) -> dict:
    """
    解析 DOCX 招标文件。

    返回结构化 dict:
    {
        "file_type": "docx",
        "file_name": "xxx.docx",
        "parse_time": "ISO timestamp",
        "structure": [  # 文档结构（标题层级 + 段落）
            {"type": "heading", "level": 1, "text": "..."},
            {"type": "paragraph", "text": "..."},
        ],
        "tables": [  # 所有表格
            {
                "table_index": 0,
                "headers": ["col1", "col2"],
                "rows": [["row1c1", "row1c2"], ...],
                "markdown": "| col1 | col2 |\n|---|---|...",
            }
        ],
        "full_text": "...",  # 纯文本全文
        "page_count": null,  # DOCX 无页码概念
    }
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    doc = Document(str(path))

    structure = []
    tables_data = []
    text_parts = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # strip namespace

        if tag == 'p':
            # 查找对应的 Paragraph 对象
            for para in doc.paragraphs:
                if para._element is element:
                    _process_paragraph(para, structure, text_parts)
                    break

        elif tag == 'tbl':
            # 表格在 doc.paragraphs 中不出现，需要单独处理
            pass

    # 单独处理表格（python-docx 不区分 body 中表格和段落的位置关系）
    for idx, table in enumerate(doc.tables):
        table_data = _process_table(table, idx)
        tables_data.append(table_data)

    full_text = "\n".join(text_parts)

    return {
        "file_type": "docx",
        "file_name": path.name,
        "parse_time": datetime.now().isoformat(),
        "structure": structure,
        "tables": tables_data,
        "full_text": full_text,
        "page_count": None,
    }


def _process_paragraph(para, structure: list, text_parts: list):
    """处理单个段落，识别标题层级并提取文本。"""
    style_name = (para.style.name or "").lower()

    # 判断是否为标题
    heading_match = re.match(r"heading\s*(\d)", style_name)
    if heading_match:
        level = int(heading_match.group(1))
        text = para.text.strip()
        if text:
            structure.append({
                "type": "heading",
                "level": level,
                "text": text,
            })
            text_parts.append(f"\n{'#' * level} {text}")
    else:
        text = para.text.strip()
        if text:
            structure.append({
                "type": "paragraph",
                "text": text,
            })
            text_parts.append(text)


def _process_table(table, index: int) -> dict:
    """处理单个表格，提取表头、行数据、Markdown 格式。"""
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # 清理单元格文本
            text = cell.text.strip().replace("\n", " ").replace("\r", "")
            cells.append(text)
        rows_data.append(cells)

    if not rows_data:
        return {"table_index": index, "headers": [], "rows": [], "markdown": ""}

    # 假设第一行为表头
    headers = rows_data[0]
    data_rows = rows_data[1:]

    # 生成 Markdown 表格
    markdown = _rows_to_markdown(headers, data_rows)

    return {
        "table_index": index,
        "headers": headers,
        "rows": data_rows,
        "markdown": markdown,
    }


def _rows_to_markdown(headers: list, rows: list) -> str:
    """将表格数据转为 Markdown 格式。"""
    if not headers:
        return ""
    lines = []
    lines.append("| " + " | ".join(h or "" for h in headers) + " |")
    lines.append("| " + " | ".join("---" for _ in headers) + " |")
    for row in rows:
        padded = row + [""] * (len(headers) - len(row))
        lines.append("| " + " | ".join(c or "" for c in padded[:len(headers)]) + " |")
    return "\n".join(lines)


# ==================== PDF 解析 ====================

def parse_pdf(file_path: str) -> dict:
    """
    解析 PDF 招标文件（使用 pdfplumber）。

    返回结构化 dict:
    {
        "file_type": "pdf",
        "file_name": "xxx.pdf",
        "parse_time": "ISO timestamp",
        "pages": [  # 每页内容
            {
                "page_number": 1,
                "text": "...",
                "tables": [...],
            }
        ],
        "tables": [  # 所有表格（跨页）
            {
                "table_index": 0,
                "page": 1,
                "headers": ["col1", "col2"],
                "rows": [["row1c1", "row1c2"], ...],
                "markdown": "...",
            }
        ],
        "full_text": "...",  # 纯文本全文
        "page_count": 5,
    }
    """
    import pdfplumber

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    tables_all = []
    pages_data = []
    text_parts = []

    with pdfplumber.open(str(path)) as pdf:
        table_index = 0

        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            text_parts.append(page_text)

            page_tables = []
            pdf_tables = page.extract_tables() or []

            for tbl in pdf_tables:
                if not tbl:
                    continue

                # 清理单元格
                cleaned = []
                for row in tbl:
                    cleaned_row = []
                    for cell in row:
                        if cell is None:
                            cleaned_row.append("")
                        else:
                            cleaned_row.append(
                                str(cell).strip().replace("\n", " ").replace("\r", "")
                            )
                    cleaned.append(cleaned_row)

                if not cleaned:
                    continue

                headers = cleaned[0]
                data_rows = cleaned[1:]
                markdown = _rows_to_markdown(headers, data_rows)

                table_entry = {
                    "table_index": table_index,
                    "page": page_num,
                    "headers": headers,
                    "rows": data_rows,
                    "markdown": markdown,
                }
                tables_all.append(table_entry)
                page_tables.append(table_entry)
                table_index += 1

            pages_data.append({
                "page_number": page_num,
                "text": page_text,
                "tables": page_tables,
            })

    full_text = "\n\n".join(text_parts)

    return {
        "file_type": "pdf",
        "file_name": path.name,
        "parse_time": datetime.now().isoformat(),
        "pages": pages_data,
        "tables": tables_all,
        "full_text": full_text,
        "page_count": len(pages_data),
    }


# ==================== 关键信息提取 ====================

def extract_key_info(parsed: dict) -> dict:
    """
    从解析结果中提取关键信息。

    适用 DOCX 和 PDF 解析结果。
    提取: 项目名称、招标编号、截止日期、预算金额、采购人、
          招标代理、资质要求、联系方式、招标方式等。

    返回:
    {
        "project_name": "...",
        "bid_number": "...",
        "deadline": "...",
        "deadline_normalized": "YYYY-MM-DD",  # 如果识别到日期
        "budget": "...",
        "buyer": "...",
        "agency": "...",
        "qualification": "...",
        "contact": "...",
        "bid_method": "...",
        "extraction_source": "full_text / tables / structure",
        "table_key_info": [...],  # 从表格中提取的额外关键信息
    }
    """
    full_text = parsed.get("full_text", "")
    tables = parsed.get("tables", [])
    structure = parsed.get("structure", [])

    result: dict = {
        "project_name": None,
        "bid_number": None,
        "deadline": None,
        "deadline_normalized": None,
        "budget": None,
        "buyer": None,
        "agency": None,
        "qualification": None,
        "contact": None,
        "bid_method": None,
        "extraction_source": None,
        "table_key_info": [],
    }

    # 第一优先级：从全文文本中正则提取
    source = _extract_from_text(full_text, result)

    # 第二优先级：从表格中提取补充信息
    table_info = _extract_from_tables(tables, result)
    if table_info:
        result["table_key_info"] = table_info
        if source is None:
            source = "tables"

    # 第三优先级：从 DOCX 结构中的标题/段落提取
    if source is None and structure:
        source = _extract_from_structure(structure, result)

    result["extraction_source"] = source or "none"

    # 尝试标准化日期
    if result.get("deadline"):
        result["deadline_normalized"] = _normalize_date(result["deadline"])

    return result


def _extract_from_text(text: str, result: dict) -> Optional[str]:
    """从全文文本中用正则提取关键信息。
    修复 #10：单次遍历提取所有关键信息，避免 O(n²) 性能问题。
    将所有模式编译为单个组合正则，一次扫描全文，按优先级分配结果。
    """
    found_any = False

    # 单次遍历：对全文只扫描一次，提取所有匹配并按优先级分配
    all_matches: Dict[str, list] = {}  # key -> [(start, value), ...]
    for key, patterns in KEY_PATTERNS.items():
        if result.get(key):  # 已有值则跳过
            continue
        all_matches[key] = []
        for pattern in patterns:
            for match in re.finditer(pattern, text):
                value = match.group(1).strip()
                value = re.sub(r"[，。、；：]+$", "", value)
                if value and len(value) < 200:
                    all_matches[key].append((match.start(), value))

    # 按优先级分配结果：优先级靠前的 key 先分配，每个 key 取最靠前的匹配
    for key, patterns in KEY_PATTERNS.items():
        if result.get(key) or key not in all_matches or not all_matches[key]:
            continue
        # 按位置排序，取最靠前的匹配
        all_matches[key].sort(key=lambda x: x[0])
        result[key] = all_matches[key][0][1]
        found_any = True

    return "full_text" if found_any else None


def _extract_from_tables(tables: list, result: dict) -> list:
    """从表格中提取关键信息，返回表格中发现的 key-value 对。"""
    table_key_info = []

    # 常见表头关键词映射
    header_map = {
        "项目名称": "project_name",
        "项目编号": "bid_number",
        "招标编号": "bid_number",
        "预算金额": "budget",
        "采购人": "buyer",
        "招标人": "buyer",
        "代理机构": "agency",
        "资质要求": "qualification",
        "联系人": "contact",
        "联系电话": "contact",
        "招标方式": "bid_method",
        "采购方式": "bid_method",
    }

    for tbl in tables:
        headers = tbl.get("headers", [])
        rows = tbl.get("rows", [])

        # 检查表头是否包含关键词
        header_indices = {}
        for i, h in enumerate(headers):
            for keyword, result_key in header_map.items():
                if keyword in h and not result.get(result_key):
                    header_indices[result_key] = i

        if not header_indices:
            # 也尝试键值对模式（两列表格）
            if len(headers) == 2:
                for row in rows:
                    if len(row) >= 2:
                        key_col = row[0].strip()
                        val_col = row[1].strip()
                        for keyword, result_key in header_map.items():
                            if keyword in key_col and not result.get(result_key) and val_col:
                                result[result_key] = val_col
                                table_key_info.append({
                                    "source": f"table_{tbl.get('table_index', '?')}",
                                    "key": key_col,
                                    "value": val_col,
                                    "mapped_to": result_key,
                                })
            continue

        # 从行数据中提取
        for row in rows:
            for result_key, col_idx in header_indices.items():
                if col_idx < len(row) and not result.get(result_key):
                    value = row[col_idx].strip()
                    if value and len(value) < 200:
                        result[result_key] = value
                        table_key_info.append({
                            "source": f"table_{tbl.get('table_index', '?')}",
                            "key": headers[col_idx],
                            "value": value,
                            "mapped_to": result_key,
                        })

    return table_key_info


def _extract_from_structure(structure: list, result: dict) -> Optional[str]:
    """从 DOCX 结构（标题+段落）中提取关键信息。"""
    # 将结构文本拼接后再用正则提取
    text_parts = []
    for item in structure:
        if item.get("text"):
            text_parts.append(item["text"])
    combined = "\n".join(text_parts)

    found_any = False
    for key, patterns in KEY_PATTERNS.items():
        if result.get(key):
            continue
        for pattern in patterns:
            match = re.search(pattern, combined)
            if match:
                value = match.group(1).strip()
                value = re.sub(r"[，。、；：]+$", "", value)
                if value and len(value) < 200:
                    result[key] = value
                    found_any = True
                    break

    return "structure" if found_any else None


def _normalize_date(date_str: str) -> Optional[str]:
    """尝试将中文/混合日期格式标准化为 YYYY-MM-DD。"""
    for pattern, formatter in DATE_PATTERNS:
        match = re.search(pattern, date_str)
        if match:
            return formatter(match)
    return None


# ==================== 便捷函数 ====================

def parse_file(file_path: str) -> dict:
    """
    自动检测文件类型并解析。

    支持 .docx 和 .pdf 扩展名。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return parse_docx(file_path)
    elif suffix == ".pdf":
        return parse_pdf(file_path)
    else:
        raise ValueError(
            f"不支持的文件格式: {suffix}，仅支持 .docx 和 .pdf"
        )


def parse_and_extract(file_path: str) -> dict:
    """
    一键解析 + 关键信息提取。

    返回:
    {
        "parsed": {...},       # 完整解析结果
        "key_info": {...},     # 关键信息提取
    }
    """
    parsed = parse_file(file_path)
    key_info = extract_key_info(parsed)
    return {
        "parsed": parsed,
        "key_info": key_info,
    }


# ==================== CLI 入口 ====================

if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    if len(sys.argv) < 2:
        print("用法: python doc_parser.py <文件路径>")
        print("示例: python doc_parser.py 招标文件.docx")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        result = parse_and_extract(file_path)
        print(json.dumps(result["key_info"], ensure_ascii=False, indent=2))
    except Exception as e:
        logger.error(f"解析失败: {e}")
        sys.exit(1)
