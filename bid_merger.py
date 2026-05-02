"""
标书整合引擎 — Day 4 任务 1
============================
功能:
  1. 合并三套标书（报价/商务/技术）为完整文档
  2. 添加统一封面页（项目名称、投标单位、日期）
  3. 自动生成目录页
  4. 统一页眉页脚格式
  5. 页码连续编号
  6. 输出完整 .docx 文件
"""

import os
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from lxml import etree

logger = logging.getLogger(__name__)

# 输出目录
BID_OUTPUT_DIR = Path("/tmp/bid-outputs")
BID_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# 模板目录
TEMPLATES_DIR = Path(__file__).parent / "templates"


# ==================== 样式工具 ====================

def _set_font(run, name='宋体', size=12, bold=False, color=None):
    """统一设置字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run._element.rPr.rFonts.set(qn('w:ascii'), name)


def _add_styled_heading(doc, text, level=1, font_name='黑体', font_size=None,
                         space_before=12, space_after=6):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        fs = font_size if font_size else ({1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12))
        _set_font(run, name=font_name, size=fs, bold=True)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    return heading


def _add_para_with_style(doc, text='', font_size=12, bold=False,
                          space_before=None, space_after=None,
                          first_line_indent=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          color=None):
    """添加带样式的段落"""
    para = doc.add_paragraph(text)
    para.alignment = alignment
    for run in para.runs:
        _set_font(run, size=font_size, bold=bold, color=color)
    if space_before is not None:
        para.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
    return para


def _add_kv_para(doc, key, value, font_size=11):
    """添加键值对段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(1)
    run_key = para.add_run(f'{key}：')
    _set_font(run_key, size=font_size, bold=True)
    run_val = para.add_run(str(value))
    _set_font(run_val, size=font_size)
    return para


def _add_bullet(doc, text, level=0, font_size=11):
    """添加列表项"""
    para = doc.add_paragraph()
    indent = 1 + level * 1.5
    para.paragraph_format.left_indent = Cm(indent)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    bullet = '•' if level == 0 else '◦'
    run_bullet = para.add_run(f'{bullet} ')
    _set_font(run_bullet, size=font_size, bold=(level == 0))
    run_text = para.add_run(text)
    _set_font(run_text, size=font_size)
    return para


def _add_simple_table(doc, headers, rows, col_widths=None):
    """添加简单表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        _set_font(run, size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shd = etree.SubElement(shading, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'D9E2F3')
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            _set_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table


# ==================== 封面生成 ====================

def _add_cover_page(doc, project_name: str, company_name: str,
                     project_type: str = "物联网卡"):
    """添加封面页"""
    now_str = datetime.now().strftime('%Y年%m月%d日')

    # 空行
    for _ in range(4):
        doc.add_paragraph()

    # 主标题
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(60)
    cover_title.paragraph_format.space_after = Pt(30)
    run_title = cover_title.add_run('投 标 文 件')
    _set_font(run_title, name='黑体', size=36, bold=True, color=(0, 51, 102))

    # 副标题
    cover_subtitle = doc.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_subtitle.paragraph_format.space_after = Pt(40)
    run_sub = cover_subtitle.add_run(f'（{project_type}项目）')
    _set_font(run_sub, name='黑体', size=22, bold=False, color=(102, 102, 102))

    # 项目名称
    cover_proj = doc.add_paragraph()
    cover_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_proj.paragraph_format.space_after = Pt(50)
    run_p = cover_proj.add_run(f'项目名称：{project_name}')
    _set_font(run_p, name='楷体', size=16, bold=False)

    # 投标单位
    cover_comp = doc.add_paragraph()
    cover_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_comp.paragraph_format.space_after = Pt(50)
    run_c = cover_comp.add_run(f'投标单位：{company_name}')
    _set_font(run_c, name='楷体', size=16, bold=False)

    # 日期
    cover_date = doc.add_paragraph()
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_date.paragraph_format.space_after = Pt(20)
    run_d = cover_date.add_run(now_str)
    _set_font(run_d, name='楷体', size=14, bold=False)

    doc.add_page_break()


# ==================== 目录生成 ====================

def _add_toc_page(doc, chapters: List[tuple]):
    """
    添加目录页
    chapters: [(章节编号, 章节标题), ...]
    """
    _add_styled_heading(doc, '目    录', level=1, font_size=18)
    doc.add_paragraph()

    for num, title_text in chapters:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_before = Pt(4)
        toc_para.paragraph_format.space_after = Pt(4)
        toc_para.paragraph_format.tab_stops.add_tab_stop(Cm(15))
        run_num = toc_para.add_run(f'{num}  ')
        _set_font(run_num, size=12, bold=True)
        # 添加 tab 和点线
        run_tab = toc_para.add_run('\t')
        _set_font(run_tab, size=12)
        run_text = toc_para.add_run(title_text)
        _set_font(run_text, size=12)

    doc.add_page_break()


# ==================== 各章节生成 ====================

def _add_bid_letter(doc, project_name: str, company_name: str,
                     contact_person: str, contact_phone: str,
                     bid_amount: float):
    """第一章：投标函"""
    now_str = datetime.now().strftime('%Y年%m月%d日')
    _add_styled_heading(doc, '第一章  投标函', level=1)
    doc.add_paragraph()

    _add_para_with_style(doc, f'致：{project_name}招标方', font_size=12,
                          space_before=12, space_after=6)
    _add_para_with_style(doc,
        f'    根据贵方发布的招标文件，{company_name}（以下简称"我方"）经认真研究招标文件的全部内容后，'
        f'决定参与本项目的投标。现正式提交投标文件，具体内容如下：',
        font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    一、我方完全理解并接受招标文件的全部内容和要求，愿意按照招标文件的规定提供相关产品与服务。',
        font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    二、我方投标总报价为人民币（大写）：¥{bid_amount:,.2f}（含税）。'
        f'该报价包含设备费、运输费、安装调试费、培训费、税费及售后服务等全部费用。',
        font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    三、我方承诺在中标后按照招标文件及合同约定履行全部责任和义务，保证按期、按质完成项目交付。',
        font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    四、我方投标文件自投标截止日起有效期为 90 个日历日。',
        font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    五、如我方中标，我方承诺在合同签订后按照约定时间完成项目交付，并提供完善的售后服务。',
        font_size=12, space_before=6, space_after=6)

    doc.add_paragraph()
    _add_para_with_style(doc, f'投标单位（盖章）：{company_name}', font_size=12, space_before=12, space_after=6)
    _add_para_with_style(doc, f'法定代表人或授权代表（签字）：{contact_person}', font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc, f'联系电话：{contact_phone}', font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc, f'日期：{now_str}', font_size=12, space_before=6, space_after=6)


def _add_auth_letter(doc, company_name: str, contact_person: str,
                      contact_phone: str, project_name: str, location: str):
    """第二章：法定代表人授权书"""
    now_str = datetime.now().strftime('%Y年%m月%d日')
    _add_styled_heading(doc, '第二章  法定代表人授权书', level=1)
    doc.add_paragraph()

    _add_para_with_style(doc,
        f'    本授权书声明：{company_name}的法定代表人（姓名、职务）在此授权{contact_person}（姓名、职务）'
        f'为我方合法代理人，以我方名义参加"{project_name}"项目的投标活动。'
        f'代理人在投标过程中签署的一切文件和处理与之有关的一切事务，我方均予以承认。',
        font_size=12, space_before=6, space_after=6)

    doc.add_paragraph()
    _add_kv_para(doc, '授权单位', company_name)
    _add_kv_para(doc, '被授权人', contact_person)
    _add_kv_para(doc, '联系电话', contact_phone)
    _add_kv_para(doc, '项目名称', project_name)
    _add_kv_para(doc, '项目地点', location)
    _add_kv_para(doc, '授权日期', now_str)

    doc.add_paragraph()
    _add_para_with_style(doc,
        '法定代表人签字：_____________    被授权人签字：_____________',
        font_size=12, space_before=12, space_after=6)
    doc.add_page_break()


def _add_price_table(doc, project_name: str, project_id: int,
                      budget: float, bid_amount: float,
                      project_type: str = "物联网卡",
                      custom_fields: Optional[Dict] = None):
    """第三章：投标报价一览表"""
    _add_styled_heading(doc, '第三章  投标报价一览表', level=1)
    doc.add_paragraph()

    _add_kv_para(doc, '项目名称', project_name)
    _add_kv_para(doc, '项目编号', f'BID-{project_id}-{datetime.now().strftime("%Y%m%d")}')
    if budget and budget > 0:
        _add_kv_para(doc, '项目预算', f'¥{budget:,.2f}')
    _add_kv_para(doc, '投标报价', f'¥{bid_amount:,.2f}')

    if budget and budget > 0:
        ratio = (bid_amount / budget * 100)
        _add_kv_para(doc, '报价占预算比例', f'{ratio:.1f}%')

    doc.add_paragraph()
    _add_styled_heading(doc, '报价明细表', level=3, font_size=13)
    doc.add_paragraph()

    cf = custom_fields or {}
    if project_type == "物联网卡":
        qty = int(cf.get('delivery_quantity', 1000))
        unit_price = bid_amount / qty if qty > 0 else bid_amount
        headers = ['序号', '费用项目', '单价（元）', '数量', '金额（元）', '备注']
        rows = [
            ['1', '物联网卡（SIM卡）', f'{unit_price * 0.4:.2f}', f'{qty}', f'{bid_amount * 0.4:.2f}', cf.get('card_type', '标准SIM卡')],
            ['2', '流量套餐费', f'{unit_price * 0.25:.2f}', f'{qty}', f'{bid_amount * 0.25:.2f}', cf.get('data_plan', '月流量1GB')],
            ['3', '平台接入费', f'{bid_amount * 0.1:.2f}', '1', f'{bid_amount * 0.1:.2f}', '卡管理平台'],
            ['4', 'API对接费', f'{bid_amount * 0.05:.2f}', '1', f'{bid_amount * 0.05:.2f}', '接口开发'],
            ['5', '技术支持服务费', f'{bid_amount * 0.1:.2f}', '1', f'{bid_amount * 0.1:.2f}', '含售后'],
            ['6', '运输及杂费', f'{bid_amount * 0.05:.2f}', '1', f'{bid_amount * 0.05:.2f}', '物流配送'],
            ['合计', '', '', '', f'{bid_amount:.2f}', '含税'],
        ]
    else:
        qty = int(cf.get('delivery_quantity', 50))
        unit_price = bid_amount / qty if qty > 0 else bid_amount
        headers = ['序号', '费用项目', '单价（元）', '数量', '金额（元）', '备注']
        rows = [
            ['1', '布控球设备', f'{unit_price * 0.5:.2f}', f'{qty}', f'{bid_amount * 0.5:.2f}', cf.get('video_resolution', '4K超清')],
            ['2', '安装材料及配件', f'{unit_price * 0.1:.2f}', f'{qty}', f'{bid_amount * 0.1:.2f}', '支架、线缆等'],
            ['3', '云存储服务费', f'{bid_amount * 0.1:.2f}', '1', f'{bid_amount * 0.1:.2f}', cf.get('storage', '云存储')],
            ['4', '安装调试费', f'{bid_amount * 0.1:.2f}', '1', f'{bid_amount * 0.1:.2f}', '现场实施'],
            ['5', '智能分析模块', f'{bid_amount * 0.1:.2f}', '1', f'{bid_amount * 0.1:.2f}', 'AI算法'],
            ['6', '培训及运维费', f'{bid_amount * 0.05:.2f}', '1', f'{bid_amount * 0.05:.2f}', '培训+运维'],
            ['合计', '', '', '', f'{bid_amount:.2f}', '含税'],
        ]

    _add_simple_table(doc, headers, rows, col_widths=[1.5, 3.5, 2.5, 2, 2.5, 3])

    doc.add_paragraph()
    _add_para_with_style(doc,
        '注：以上报价为含税总价，报价有效期为 90 个日历日，自投标截止日起计算。',
        font_size=10, space_before=6, space_after=6)


def _add_technical_section(doc, project_name: str, location: str,
                            budget: float, project_type: str = "物联网卡",
                            custom_fields: Optional[Dict] = None):
    """第四章：技术方案"""
    cf = custom_fields or {}

    _add_styled_heading(doc, '第四章  技术方案', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '4.1 项目概述', level=2, font_size=14)
    _add_para_with_style(doc,
        f'    本项目为"{project_name}"，项目地点位于{location}，'
        f'项目预算为¥{budget:,.2f}。我方凭借在该领域的丰富经验和专业技术实力，提供完整的技术解决方案。',
        font_size=12, space_before=6, space_after=6)

    _add_styled_heading(doc, '4.2 技术路线与架构', level=2, font_size=14)
    if project_type == "物联网卡":
        _add_para_with_style(doc,
            '    我方采用"终端-网络-平台"三层架构的物联网卡解决方案，确保系统稳定可靠：',
            font_size=12, space_before=6, space_after=6)
        _add_bullet(doc, '终端层：支持多类型物联网卡（SIM/eSIM/M2M），适配各类物联网设备')
        _add_bullet(doc, '网络层：支持 FDD-LTE/TDD-LTE/NB-IoT/5G 全频段覆盖，保障网络质量')
        _add_bullet(doc, '平台层：自建卡管理平台，提供卡生命周期管理、用量监控、故障诊断等核心功能')
        _add_bullet(doc, '安全层：端到端加密传输，双向认证机制，确保数据安全')
    else:
        _add_para_with_style(doc,
            '    我方采用"前端采集-网络传输-云端分析"三层架构的布控球解决方案：',
            font_size=12, space_before=6, space_after=6)
        _add_bullet(doc, '前端采集层：4K超高清布控球设备，支持多种供电方式（太阳能/市电/电池）')
        _add_bullet(doc, '网络传输层：4G/5G无线传输，支持RTSP/ONVIF/GB/T28181标准协议')
        _add_bullet(doc, '云端分析层：视频云平台+AI智能分析，支持人脸识别、车牌识别、行为分析')
        _add_bullet(doc, '安全层：视频流加密传输，权限分级管控，符合等保要求')

    _add_styled_heading(doc, '4.3 核心技术参数', level=2, font_size=14)
    doc.add_paragraph()

    if project_type == "物联网卡":
        spec_headers = ['参数项', '技术指标', '说明']
        spec_rows = [
            ['卡类型', cf.get('card_type', 'SIM卡（可定制eSIM/M2M）'), '支持多种卡类型'],
            ['频段支持', cf.get('frequency_band', '全频段（FDD-LTE/TDD-LTE/NB-IoT/5G）'), '广泛兼容性'],
            ['流量套餐', cf.get('data_plan', '月流量1GB（可定制）'), '灵活套餐选择'],
            ['运营商合作', cf.get('operator', '中国移动/联通/电信'), '多运营商冗余'],
            ['APN配置', cf.get('apn', '支持自定义APN'), '灵活网络配置'],
            ['QoS保障', cf.get('qos', '优先服务级别'), '网络质量保障'],
            ['安全加密', cf.get('security', '双向认证+端到端加密'), '高安全等级'],
        ]
    else:
        spec_headers = ['参数项', '技术指标', '说明']
        spec_rows = [
            ['视频分辨率', cf.get('video_resolution', '4K超高清（3840×2160）'), '高清画质'],
            ['帧率', cf.get('frame_rate', '30fps（支持60fps）'), '流畅视频'],
            ['码率', cf.get('bit_rate', '4Mbps-10Mbps自适应'), '智能码率控制'],
            ['存储方式', cf.get('storage', '本地+云存储双模式'), '可靠存储'],
            ['供电方式', cf.get('power', '太阳能+电池混合供电'), '无市电场景适用'],
            ['安装方式', cf.get('installation', '立杆/墙面/车载多模式'), '灵活部署'],
            ['无线传输', cf.get('wireless', '4G/5G/Wi-Fi混合组网'), '多网络冗余'],
            ['协议支持', cf.get('protocol', 'RTSP/ONVIF/GB/T28181'), '标准协议兼容'],
            ['智能分析', '人脸识别/车牌识别/行为分析', 'AI算法支持'],
        ]

    _add_simple_table(doc, spec_headers, spec_rows, col_widths=[3.5, 6, 4])
    doc.add_page_break()


def _add_implementation_plan(doc, project_type: str = "物联网卡"):
    """第五章：项目实施计划"""
    _add_styled_heading(doc, '第五章  项目实施计划', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '5.1 实施流程', level=2, font_size=14)
    _add_para_with_style(doc, '    我方将按照以下阶段推进项目实施：', font_size=12, space_before=6, space_after=6)

    implementation_headers = ['阶段', '时间', '主要工作内容', '交付物']
    if project_type == "物联网卡":
        implementation_rows = [
            ['第一阶段\n需求确认', '第1周', '需求调研、方案确认、APN配置确认', '需求确认书'],
            ['第二阶段\n卡片生产', '第2-3周', '卡片定制生产、质量检测、编号管理', '产品检测报告'],
            ['第三阶段\n平台部署', '第3-4周', '管理平台部署、API对接调试', '平台部署报告'],
            ['第四阶段\n物流交付', '第4-5周', '产品包装、物流配送、到货签收', '签收单'],
            ['第五阶段\n联调验收', '第5-6周', '现场联调测试、性能验证、项目验收', '验收报告'],
        ]
    else:
        implementation_rows = [
            ['第一阶段\n需求调研', '第1周', '现场勘测、需求确认、方案设计', '需求调研报告'],
            ['第二阶段\n设备采购', '第2-3周', '设备采购、出厂检测、配件准备', '设备检测报告'],
            ['第三阶段\n安装部署', '第3-5周', '现场安装、网络配置、平台对接', '安装部署报告'],
            ['第四阶段\n系统联调', '第5-6周', '系统联调、功能测试、性能优化', '测试报告'],
            ['第五阶段\n培训验收', '第6-7周', '用户培训、项目验收、交付文档', '验收报告'],
        ]

    _add_simple_table(doc, implementation_headers, implementation_rows, col_widths=[3, 2, 5.5, 3])

    _add_styled_heading(doc, '5.2 项目团队', level=2, font_size=14)
    _add_para_with_style(doc, '    我方将组建专业项目团队，确保项目高质量交付：', font_size=12, space_before=6, space_after=6)

    team_headers = ['角色', '人数', '职责']
    team_rows = [
        ['项目经理', '1人', '项目统筹管理、进度控制、对外沟通'],
        ['技术负责人', '1人', '技术方案设计、技术难题攻关'],
        ['实施工程师', '2-4人', '现场安装部署、系统调试'],
        ['测试工程师', '1人', '功能测试、性能测试、验收测试'],
        ['售后工程师', '1人', '售后支持、运维保障'],
    ]
    _add_simple_table(doc, team_headers, team_rows, col_widths=[3, 2, 8.5])

    _add_styled_heading(doc, '5.3 质量保障措施', level=2, font_size=14)
    _add_bullet(doc, '严格执行 ISO 9001 质量管理体系')
    _add_bullet(doc, '每阶段设立质量检查点，不合格返工')
    _add_bullet(doc, '关键节点邀请第三方检测机构参与')
    _add_bullet(doc, '建立项目日报/周报机制，及时沟通项目进展')
    doc.add_page_break()


def _add_after_sales(doc, project_type: str = "物联网卡"):
    """第六章：售后服务方案"""
    _add_styled_heading(doc, '第六章  售后服务方案', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '6.1 服务承诺', level=2, font_size=14)
    _add_bullet(doc, '质保期：自验收合格之日起 12 个月')
    _add_bullet(doc, '响应时间：7×24小时响应，2小时内远程响应，24小时内现场到达')
    _add_bullet(doc, '故障处理：一般故障 4 小时内解决，重大故障 24 小时内提供备用方案')
    _add_bullet(doc, '定期巡检：每季度一次现场巡检，提交巡检报告')
    _add_bullet(doc, '技术支持：免费技术培训 2 次/年，不定期技术交流')

    _add_styled_heading(doc, '6.2 售后服务体系', level=2, font_size=14)
    _add_para_with_style(doc, '    我方建立了完善的三级售后服务体系：', font_size=12, space_before=6, space_after=6)
    _add_bullet(doc, '一级支持：客服热线、在线工单系统、自助知识库')
    _add_bullet(doc, '二级支持：远程诊断、远程修复、远程升级')
    _add_bullet(doc, '三级支持：现场服务、备件更换、设备维修')

    _add_styled_heading(doc, '6.3 备品备件保障', level=2, font_size=14)
    _add_para_with_style(doc,
        '    我方承诺在质保期内免费提供备品备件服务。对于关键设备，'
        '我方将在项目所在地设立备件库，确保故障设备能快速替换，不影响项目运行。',
        font_size=12, space_before=6, space_after=6)

    _add_styled_heading(doc, '6.4 培训计划', level=2, font_size=14)
    training_headers = ['培训阶段', '培训内容', '培训时长', '培训对象']
    training_rows = [
        ['第一阶段', '产品功能与操作', '4学时', '操作人员'],
        ['第二阶段', '日常维护与故障排查', '4学时', '运维人员'],
        ['第三阶段', '高级功能与系统管理', '4学时', '管理人员'],
    ]
    _add_simple_table(doc, training_headers, training_rows, col_widths=[3, 5, 2.5, 3])
    doc.add_page_break()


def _add_qualifications(doc, company_name: str, project_type: str = "物联网卡"):
    """第七章：企业资质与业绩"""
    _add_styled_heading(doc, '第七章  企业资质与业绩', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '7.1 企业资质', level=2, font_size=14)

    if project_type == "物联网卡":
        qual_rows = [
            ['营业执照', '合法有效，经营范围包含物联网相关业务'],
            ['增值电信业务经营许可证', '含ICP许可证、EDI证书'],
            ['物联网相关资质', '物联网行业准入资质'],
            ['ISO 9001 质量管理体系认证', '通过国际质量体系认证'],
            ['ISO 27001 信息安全管理体系', '通过信息安全管理体系认证'],
        ]
    else:
        qual_rows = [
            ['营业执照', '合法有效，经营范围包含安防相关业务'],
            ['安防工程企业资质', '含安防设计施工维护资质'],
            ['通信工程施工资质', '通信设备安装调试资质'],
            ['信息系统集成资质', '系统集成能力认证'],
            ['ISO 9001 质量管理体系认证', '通过国际质量体系认证'],
        ]
    _add_simple_table(doc, ['资质名称', '资质说明'], qual_rows, col_widths=[5, 8.5])

    _add_styled_heading(doc, '7.2 类似项目业绩', level=2, font_size=14)
    _add_para_with_style(doc,
        f'    {company_name}在{project_type}领域积累了丰富的项目经验，以下为部分代表性案例：',
        font_size=12, space_before=6, space_after=6)

    if project_type == "物联网卡":
        case_rows = [
            ['某省电力物联网卡项目', '50000张', '物联网卡供应+平台接入', '2025年'],
            ['某市智慧交通物联网卡项目', '30000张', '物联网卡+API对接', '2025年'],
            ['某物流企业物联网卡项目', '20000张', '物联网卡+管理平台', '2024年'],
            ['某农业物联网监测项目', '10000张', 'NB-IoT卡+云平台', '2024年'],
        ]
    else:
        case_rows = [
            ['某市雪亮工程布控球项目', '200台', '设备供应+安装+平台', '2025年'],
            ['某景区智能监控布控球项目', '80台', '设备+AI分析+云存储', '2025年'],
            ['某工地安全监控布控球项目', '150台', '设备供应+安装运维', '2024年'],
            ['某河道防洪监控布控球项目', '60台', '太阳能布控球+云平台', '2024年'],
        ]
    _add_simple_table(doc, ['项目名称', '项目规模', '服务内容', '完成时间'], case_rows, col_widths=[4.5, 2.5, 4.5, 2])
    doc.add_page_break()


def _add_project_analysis(doc, project_name: str, description: str,
                           project_type: str = "物联网卡"):
    """第八章：项目理解与需求分析"""
    _add_styled_heading(doc, '第八章  项目理解与需求分析', level=1)
    doc.add_paragraph()

    _add_styled_heading(doc, '8.1 项目背景理解', level=2, font_size=14)
    _add_para_with_style(doc,
        f'    通过对"{project_name}"项目的深入研究，我方理解该项目的主要背景为：',
        font_size=12, space_before=6, space_after=6)
    _add_para_with_style(doc,
        f'    {description or "本项目旨在通过采购相关设备与服务，提升业务能力和管理效率。"}',
        font_size=12, space_before=6, space_after=6)

    _add_styled_heading(doc, '8.2 核心需求分析', level=2, font_size=14)
    if project_type == "物联网卡":
        _add_bullet(doc, '稳定可靠的网络连接：保障物联网设备全天候在线')
        _add_bullet(doc, '灵活的流量管理：根据业务需求提供差异化流量方案')
        _add_bullet(doc, '完善的平台支撑：实现卡的集中管理和远程运维')
        _add_bullet(doc, '安全保障机制：确保数据传输和存储安全')
        _add_bullet(doc, '快速响应服务：建立高效的故障处理机制')
    else:
        _add_bullet(doc, '高质量视频采集：确保监控画面清晰流畅')
        _add_bullet(doc, '可靠的传输网络：保障视频数据实时传输')
        _add_bullet(doc, '智能分析能力：利用AI技术提升监控效率')
        _add_bullet(doc, '灵活的部署方案：适应不同场景的安装需求')
        _add_bullet(doc, '完善的售后服务：保障系统长期稳定运行')

    _add_styled_heading(doc, '8.3 我方优势', level=2, font_size=14)
    _add_bullet(doc, '行业经验丰富：在相关领域拥有多年项目经验')
    _add_bullet(doc, '技术实力雄厚：自有研发团队，持续技术创新')
    _add_bullet(doc, '服务网络完善：覆盖全国的服务网点')
    _add_bullet(doc, '成本优势明显：规模采购带来的价格优势')
    _add_bullet(doc, '质量保证体系：通过多项国际认证')


def _add_appendix(doc, project_type: str = "物联网卡"):
    """附录"""
    doc.add_page_break()
    _add_styled_heading(doc, '附    录', level=1)
    doc.add_paragraph()
    _add_kv_para(doc, '文件生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    _add_kv_para(doc, '生成系统', '投标公司赚钱引擎 AI 标书生成系统 v2.0')
    _add_kv_para(doc, '项目类型', project_type)
    _add_kv_para(doc, '文档版本', 'V1.0')
    _add_kv_para(doc, '文档说明', '本文档由系统自动生成，仅供参考，正式投标前请由专业人员审核完善。')


# ==================== 页眉页脚设置 ====================

def _setup_header_footer(doc, project_name: str, company_name: str):
    """为所有 section 设置统一的页眉页脚"""
    now_str = datetime.now().strftime('%Y-%m-%d')
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 左侧：公司名
        run_left = header_para.add_run(company_name)
        _set_font(run_left, name='宋体', size=8, color=(102, 102, 102))

        # 中间分隔
        run_sep = header_para.add_run('  |  ')
        _set_font(run_sep, name='宋体', size=8, color=(102, 102, 102))

        # 右侧：项目名
        run_right = header_para.add_run(project_name[:30])
        _set_font(run_right, name='宋体', size=8, color=(102, 102, 102))

        # 页眉下划线
        header_para2 = header.add_paragraph()
        header_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_line = header_para2.add_run('─' * 80)
        _set_font(run_line, name='宋体', size=6, color=(180, 180, 180))

        # 页脚 - 页码
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_pg = footer_para.add_run('第 ')
        _set_font(run_pg, name='宋体', size=9, color=(102, 102, 102))

        # 添加 PAGE 域
        from docx.oxml.ns import qn
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_pg._element.append(fld_char_begin)

        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_pg._element.append(instr)

        fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run_pg._element.append(fld_char_sep)

        run_pg2 = footer_para.add_run('1')
        _set_font(run_pg2, name='宋体', size=9, color=(102, 102, 102))

        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_pg2._element.append(fld_char_end)

        run_of = footer_para.add_run(' 页 / 共 1 页')
        _set_font(run_of, name='宋体', size=9, color=(102, 102, 102))


# ==================== 主合并函数 ====================

def merge_bid_documents(project_id: int, output_path: str,
                         company_name: str = "默认投标公司",
                         contact_person: str = "张三",
                         contact_phone: str = "13800138000",
                         bid_amount: float = 100000.0,
                         project_type: str = "物联网卡",
                         custom_fields: Optional[Dict[str, Any]] = None,
                         get_db_connection_func=None) -> Dict[str, Any]:
    """
    合并三套标书为完整文档，返回文件路径和页数。

    Args:
        project_id: 项目 ID
        output_path: 输出文件路径
        company_name: 投标单位名称
        contact_person: 联系人
        contact_phone: 联系电话
        bid_amount: 投标金额
        project_type: 项目类型（物联网卡/布控球）
        custom_fields: 自定义字段
        get_db_connection_func: 数据库连接函数（可选）

    Returns:
        {
            "status": "success",
            "file_path": "/path/to/output.docx",
            "page_count": 20,
            "project_id": 1,
            "project_name": "xxx",
            "generated_at": "2026-05-02T19:00:00",
        }
    """
    logger.info(f"[标书整合] 开始合并项目 {project_id} 的标书...")

    # 确保输出目录存在
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # 获取项目名称和描述（如果有数据库连接）
    project_name = f"项目-{project_id}"
    project_location = "待定"
    project_budget = 0.0
    project_description = ""

    if get_db_connection_func:
        try:
            from bid_template_engine import load_templates_from_db
            templates = load_templates_from_db(project_id, get_db_connection_func)

            if templates:
                # 从模板中获取项目信息
                content = templates[0].get("content", {})
                proj_info = content.get("project_info", {})
                if proj_info.get("title"):
                    project_name = proj_info["title"]
        except Exception as e:
            logger.warning(f"[标书整合] 无法从数据库加载模板信息: {e}")

        # 尝试从 bid_notices 获取项目详情
        try:
            with get_db_connection_func() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT title, budget, region, description FROM bid_notices WHERE id = %s",
                    (project_id,)
                )
                row = cursor.fetchone()
                if row:
                    project_name = row[0] or project_name
                    project_budget = row[1] or 0.0
                    project_location = row[2] or "待定"
                    project_description = row[3] or ""
                cursor.close()
        except Exception as e:
            logger.warning(f"[标书整合] 无法从 bid_notices 获取项目详情: {e}")

    # ==================== 构建文档 ====================
    doc = Document()

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ==================== 1. 封面页 ====================
    _add_cover_page(doc, project_name, company_name, project_type)

    # ==================== 2. 目录页 ====================
    chapters = [
        ('第一章', '投标函'),
        ('第二章', '法定代表人授权书'),
        ('第三章', '投标报价一览表'),
        ('第四章', '技术方案'),
        ('第五章', '项目实施计划'),
        ('第六章', '售后服务方案'),
        ('第七章', '企业资质与业绩'),
        ('第八章', '项目理解与需求分析'),
        ('附录', '文档信息'),
    ]
    _add_toc_page(doc, chapters)

    # ==================== 3. 各章节内容 ====================
    # 第一章：投标函
    _add_bid_letter(doc, project_name, company_name, contact_person,
                     contact_phone, bid_amount)

    # 第二章：授权书
    _add_auth_letter(doc, company_name, contact_person, contact_phone,
                      project_name, project_location)

    # 第三章：报价表
    _add_price_table(doc, project_name, project_id, project_budget,
                      bid_amount, project_type, custom_fields)
    doc.add_page_break()

    # 第四章：技术方案
    _add_technical_section(doc, project_name, project_location, project_budget,
                            project_type, custom_fields)

    # 第五章：实施计划
    _add_implementation_plan(doc, project_type)

    # 第六章：售后服务
    _add_after_sales(doc, project_type)

    # 第七章：企业资质
    _add_qualifications(doc, company_name, project_type)

    # 第八章：项目理解
    _add_project_analysis(doc, project_name, project_description, project_type)

    # 附录
    _add_appendix(doc, project_type)

    # ==================== 4. 页眉页脚 ====================
    _setup_header_footer(doc, project_name, company_name)

    # ==================== 5. 保存 ====================
    doc.save(output_path)

    # 估算页数（每个 page_break + 1页）
    page_count = doc.element.xpath('.//w:br[@w:type="page"]').__len__() + 1

    logger.info(f"[标书整合] 标书已生成: {output_path} (约 {page_count} 页)")

    return {
        "status": "success",
        "file_path": str(output_path),
        "page_count": page_count,
        "project_id": project_id,
        "project_name": project_name,
        "file_size_bytes": Path(output_path).stat().st_size,
        "generated_at": datetime.now().isoformat(),
    }


# ==================== 便捷函数 ====================

def merge_bid_documents_to_default(project_id: int,
                                     company_name: str = "默认投标公司",
                                     contact_person: str = "张三",
                                     contact_phone: str = "13800138000",
                                     bid_amount: float = 100000.0,
                                     project_type: str = "物联网卡",
                                     custom_fields: Optional[Dict[str, Any]] = None,
                                     get_db_connection_func=None) -> Dict[str, Any]:
    """使用默认输出路径合并标书"""
    output_dir = BID_OUTPUT_DIR / str(project_id)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"完整标书_{project_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    output_path = str(output_dir / filename)

    return merge_bid_documents(
        project_id=project_id,
        output_path=output_path,
        company_name=company_name,
        contact_person=contact_person,
        contact_phone=contact_phone,
        bid_amount=bid_amount,
        project_type=project_type,
        custom_fields=custom_fields,
        get_db_connection_func=get_db_connection_func,
    )
